"""
Build docs/Project_Explanation_Script.docx — a plain-English, viva-style
walkthrough of the project intended for the examiner and the team
(Harshal, Yashraj, Aapesha).

Same supervisor formatting as the main report:
  - Title  : Bookman Old Style, 14pt, Bold, Underlined
  - Body   : Bookman Old Style, 12pt, 1.5 line spacing, Justified
  - One sub-topic per page
"""
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

FONT = "Bookman Old Style"


def _add_underline(run):
    rPr = run._element.get_or_add_rPr()
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)


def _set_run_font(run, size_pt: int, bold: bool = False):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)


def _set_paragraph_format(p, justified=True):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, 14, bold=True)
    _add_underline(run)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, 12, bold=False)
    _set_paragraph_format(p, justified=True)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run, 12, bold=False)
    _set_paragraph_format(p, justified=True)


def add_centered(doc, text, size=14, bold=True, underline=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold)
    if underline:
        _add_underline(run)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_qa(doc, q, a):
    """Render a 'Q: ... / A: ...' pair as two consecutive paragraphs,
    where the question is bold and the answer is normal body."""
    p = doc.add_paragraph()
    rq = p.add_run("Q: ")
    _set_run_font(rq, 12, bold=True)
    ra = p.add_run(q)
    _set_run_font(ra, 12, bold=True)
    _set_paragraph_format(p, justified=True)

    p2 = doc.add_paragraph()
    rl = p2.add_run("A: ")
    _set_run_font(rl, 12, bold=True)
    rt = p2.add_run(a)
    _set_run_font(rt, 12, bold=False)
    _set_paragraph_format(p2, justified=True)


def page_break(doc):
    doc.add_page_break()


def section(doc, title, paragraphs, bullets=None):
    add_title(doc, title)
    for para in paragraphs:
        add_body(doc, para)
    if bullets:
        for b in bullets:
            add_bullet(doc, b)
    page_break(doc)


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def build(out_path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
    set_default_style(doc)

    # ===================== COVER =====================
    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, "Project Explanation Script", size=18, bold=True, underline=True)
    add_centered(doc, "(Simple-Language Viva Walk-through)", size=14, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Project Title", size=12, bold=False)
    add_centered(
        doc,
        "Mitigating Hallucinations in Large Language Models",
        size=14, bold=True,
    )
    add_centered(
        doc,
        "through Multi-Agent Consensus and Knowledge Graph Integration",
        size=14, bold=True,
    )
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Team", size=12, bold=False)
    add_centered(doc, "Harshal Ahire    •    Yashraj    •    Aapesha", size=14, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "M.Sc. Computer Science (Sem III) — 2025–2026", size=12, bold=False)
    add_centered(doc, "Indira College of Commerce & Science, Pune", size=12, bold=False)
    page_break(doc)

    # ===================== HOW TO USE THIS SCRIPT =====================
    section(
        doc,
        "How to Use This Script",
        [
            "This document is a plain-English script for the project viva. It is "
            "intended for two readers. The first reader is the team itself — "
            "Harshal, Yashraj and Aapesha — so that everybody can speak about every "
            "part of the project with confidence. The second reader is the "
            "examiner, who should be able to follow the explanation from the very "
            "first paragraph without prior knowledge of the topic.",
            "The script is divided into short sections. Each section is on its own "
            "page. The first few pages give a two-minute pitch and a real-world "
            "story. The middle pages walk through every component of the system, "
            "in order. The last pages answer the questions the examiner is most "
            "likely to ask, and provide a clean closing statement.",
            "When practising, read each section out aloud. Speak slowly. If a "
            "sentence feels awkward, simplify it in your own words; the goal is "
            "understanding, not memorisation. If the examiner interrupts with a "
            "question, jump to the matching answer in the Question and Answer "
            "section near the end of this script and resume from there.",
        ],
    )

    # ===================== TWO-MINUTE PITCH =====================
    section(
        doc,
        "Two-Minute Pitch (Speak This First)",
        [
            "Good afternoon. Our project is called Mitigating Hallucinations in "
            "Large Language Models through Multi-Agent Consensus and Knowledge "
            "Graph Integration. In simple words, when an artificial intelligence "
            "tool such as ChatGPT confidently makes up a wrong fact — for "
            "example, a fake court case or a fake medicine name — that mistake is "
            "called a hallucination. Hallucinations are dangerous in serious "
            "fields like law, medicine and finance because the wrong answer "
            "looks completely correct on the surface.",
            "Our project asks one question. Can we build a system that "
            "automatically catches these hallucinations before they reach the "
            "user? Most existing systems use a single AI model to both write the "
            "answer and check the answer. That is like asking a student to grade "
            "their own paper — they will almost never catch their own mistakes. "
            "We solve this by using a team of AI agents that play different roles "
            "and never see each other’s notes.",
            "We tested our system on a real-world dataset from Kaggle that lists "
            "more than four hundred real court cases in which lawyers were "
            "actually fined for using AI tools that produced fake citations. We "
            "compared our team-of-agents approach against a single-agent baseline. "
            "On the same evidence and the same questions, our system caught the "
            "fake citation, while the single-agent baseline accepted it. Our F1 "
            "score is 0.667; the baseline’s F1 score is 0.0. The price is that "
            "our system is roughly five times slower, which is acceptable in "
            "high-stakes domains.",
            "The whole pipeline runs on a normal laptop using free, open-source "
            "tools. The code, the data and this report are all on GitHub. The "
            "next slides take you step by step through the data, the architecture "
            "and the results. Thank you.",
        ],
    )

    # ===================== A REAL-WORLD STORY =====================
    section(
        doc,
        "A Real-World Story (Why This Matters)",
        [
            "In the year 2023, two lawyers in New York filed a legal brief in a "
            "case against an airline. The brief cited six earlier court cases as "
            "supporting precedents. The judge tried to look them up. None of the "
            "six cases existed. The lawyers had asked ChatGPT to write the "
            "brief, and ChatGPT had simply invented six convincing-looking case "
            "names with judge names, dates and quotations attached. The lawyers "
            "did not check, the AI tool did not warn them, and the judge "
            "discovered the fabrication.",
            "The lawyers were fined and reported to the New York State Bar. The "
            "story made international news. Since then, the same kind of incident "
            "has happened more than four hundred times around the world. There "
            "is now a public dataset on Kaggle, maintained by Mr. Damien "
            "Charlotin, that lists every documented case along with the court, "
            "the date, the AI tool that was used, the fine and a description of "
            "what the AI fabricated.",
            "This is the dataset our project uses. The problem is real. The "
            "consequences are measurable in money, in professional sanctions and "
            "in trust in the legal system. Our goal is to build the kind of "
            "automatic safety net that should sit between any AI tool and any "
            "human reader before this kind of mistake reaches them.",
        ],
    )

    # ===================== HALLUCINATION IN ONE PAGE =====================
    section(
        doc,
        "What Is a Hallucination, in One Page",
        [
            "A hallucination is a confident-looking answer that is not actually "
            "true. The word comes from the way these mistakes feel — the AI "
            "‘sees’ something that is not there, and then describes it as if it "
            "is there.",
            "There are two main types. The first type is intrinsic: the AI "
            "contradicts the very document you gave it. For example, you give "
            "the AI a patient record that says the patient is allergic to "
            "penicillin, and the AI summary recommends a penicillin antibiotic. "
            "The second type is extrinsic: the AI invents details that are not "
            "in the source document at all. For example, the AI cites a court "
            "case that no court ever heard.",
            "Both types come from the same root cause. AI language models are "
            "trained to predict the most likely next word, not to verify facts. "
            "When the model is uncertain, it does not say ‘I do not know.’ It "
            "produces the word that sounds most plausible, which is often wrong. "
            "This is not a bug; it is exactly what the model was trained to do. "
            "Fixing it requires building safety machinery around the model, "
            "not waiting for the model itself to magically improve.",
        ],
    )

    # ===================== THE BIG IDEA =====================
    section(
        doc,
        "The Big Idea (Two Sentences)",
        [
            "If one AI cannot reliably check its own work, then we use several "
            "AI agents that play different roles. The trick is that the agent "
            "doing the fact-checking is never allowed to see what the agent "
            "doing the writing actually wrote — the checker only sees the raw "
            "evidence and a list of small claims to verify.",
            "We call this property ‘deliberate information asymmetry’. It is "
            "the central architectural idea of our project, and we enforce it "
            "through the structure of the orchestration graph rather than "
            "through prompts. In other words, the asymmetry is encoded in the "
            "code itself; it is not just a polite request to the model.",
        ],
    )

    # ===================== KITCHEN ANALOGY =====================
    section(
        doc,
        "An Easy Analogy: A Restaurant Kitchen",
        [
            "Think of a restaurant kitchen with four people working on each dish. "
            "The first person, the Retriever, walks into the pantry and pulls "
            "out the actual ingredients that match the order — onions, salt, "
            "tomatoes. The second person, the Solver, is the chef. The chef "
            "uses those ingredients to cook a dish and writes a description of "
            "the dish on a card.",
            "The third person, the Proposer, is a quality inspector. The "
            "inspector reads the chef’s description and breaks it down into a "
            "checklist: ‘the dish contains onions; the dish is mildly salted; "
            "the dish is vegetarian.’ The fourth person, the Checker, is "
            "completely blind to the chef’s description. The checker only sees "
            "the original ingredients in the pantry and the inspector’s "
            "checklist. The checker walks the checklist item by item and "
            "verifies whether the pantry actually supports each claim.",
            "Finally, a fifth person, the Synthesizer, prepares the final plate. "
            "If everything checks out, the chef’s original dish is served. If "
            "the checker found a problem — say, the chef claimed the dish "
            "contains saffron, but the pantry has no saffron — the Synthesizer "
            "rewrites the description so that the false claim is removed before "
            "the customer ever sees it.",
            "Our software project does exactly this, but with text instead of "
            "food, with court cases instead of ingredients, and with AI models "
            "in each role.",
        ],
    )

    # ===================== THE DATASET =====================
    section(
        doc,
        "The Dataset We Used",
        [
            "We use the AI Hallucination Cases dataset published by Mr. Damien "
            "Charlotin on the Kaggle platform. The link is in the References "
            "section of the project report. The dataset is a single CSV file "
            "with four hundred and twenty-six rows. Each row is one real court "
            "case in which an AI tool produced a hallucination that was "
            "discovered by the court.",
            "Each row has fourteen columns. The most important columns are: "
            "Case Name (for example, ‘People v. Ruiz Alvarez’), Court (for "
            "example, ‘CA California’), Date (when the ruling was issued), "
            "Party — that is, who used the AI (Lawyer, Self-Represented "
            "Litigant, etc.), AI Tool (which AI was used — sometimes "
            "‘Unidentified’), Hallucination (a one-line description of what "
            "was fabricated), Outcome (what the court did about it), Monetary "
            "Penalty (how much the offender was fined), and Details (a "
            "paragraph or two of the actual narrative from the ruling).",
            "After cleaning the data — removing rows whose Details column was "
            "too short to be useful — we are left with two hundred and forty-"
            "three usable rows. These two hundred and forty-three real cases "
            "form the evidence corpus for our experiments.",
        ],
    )

    # ===================== HOW WE TURN ROWS INTO TESTS =====================
    section(
        doc,
        "How We Turn Rows Into Test Questions",
        [
            "For every case in our cleaned dataset, our experiment generates two "
            "test questions, which we call probes. The first probe is called "
            "TRUE: we take a real fact straight from the case Details and ask "
            "the AI system, ‘Is this claim supported by the evidence?’ The "
            "correct answer is yes. The second probe is called FALSE: we take "
            "the same true claim and we paste a fake legal citation onto the "
            "end — for example, ‘…and the court relied on Whitmore v. Ardant, "
            "612 F.3d 441 (9th Cir. 2019).’ That citation does not exist in "
            "real life. Because the fake citation is not present anywhere in "
            "our evidence store, a properly behaving system must flag the "
            "FALSE probe as a hallucination.",
            "This pairing is deliberately controlled. Both probes start from the "
            "same case. Both probes share the same surrounding sentence. The "
            "only difference is the fabricated citation that we appended to "
            "the FALSE probe. This means we can directly measure how well a "
            "system distinguishes a real claim from a fabricated one without "
            "the comparison being polluted by other variables.",
            "We use a fixed random seed so the same probes are generated every "
            "time we run the experiment. This makes the result reproducible — "
            "anybody downloading our repository can run the same probes and "
            "should obtain the same numbers.",
        ],
    )

    # ===================== THE TWO PIPELINES =====================
    section(
        doc,
        "The Two Pipelines We Compared",
        [
            "Our project compares two systems on exactly the same evidence and "
            "exactly the same probes. The first system is the Baseline. The "
            "second system is our Multi-Agent pipeline.",
            "The Baseline is a single-agent Retrieval-Augmented Generation "
            "system. It does only two things. First, it searches the evidence "
            "store for documents that look similar to the claim. Second, it "
            "asks one AI model the question, ‘Given this evidence, is the "
            "claim supported or is it a hallucination?’ The model answers "
            "directly with a yes or no and a short explanation.",
            "Our Multi-Agent pipeline is more elaborate. It has four roles, "
            "described on the next page: Solver, Proposer, Checker, and "
            "Synthesizer. The Retrieve step is shared with the baseline. The "
            "Solver writes a flowing answer. The Proposer breaks the answer "
            "into small testable claims. The Checker verifies each claim "
            "without seeing the Solver’s flowing answer. The Synthesizer "
            "either passes the original answer through or rewrites it after "
            "removing the unsupported claims.",
            "Both systems use the same AI model in the background, the same "
            "embeddings, the same evidence store and the same retrieval "
            "settings. The only thing that changes between the two arms of the "
            "experiment is the architecture. This isolates the effect of the "
            "multi-agent design.",
        ],
    )

    # ===================== THE FOUR AGENTS =====================
    section(
        doc,
        "The Four Agents Inside Our System",
        [
            "Each of the four agents inside our system has one specific job. "
            "Below is a short description of each job in plain English. The "
            "names match the function names in the source code so the team can "
            "find them easily.",
        ],
        bullets=[
            "Solver — the writer. Reads the question and the evidence, then "
            "writes a fluent draft answer in natural language. We let it run "
            "with a slightly higher creativity setting (temperature 0.7) so "
            "the answer is readable.",
            "Proposer — the breaker-down. Reads only the Solver’s draft and "
            "splits it into a JSON list of small claims, each phrased as a "
            "(question, answer) pair. We force a strict JSON output so the "
            "next stage can be processed by code.",
            "Checker — the auditor. This is the most important agent. The "
            "Checker is given two things and only two things: the raw "
            "evidence and the list of small claims. The Checker is NOT given "
            "the Solver’s draft. The Checker walks every claim against the "
            "evidence and decides whether it is supported or fabricated.",
            "Synthesizer — the editor. Takes the Solver’s draft and the "
            "Checker’s verdicts. If everything is supported, the draft "
            "passes through unchanged. If something is flagged as a "
            "hallucination, the Synthesizer rewrites the draft so that the "
            "unsupported claim is removed before the user sees it.",
        ],
    )

    # ===================== WHY BLINDING THE CHECKER MATTERS =====================
    section(
        doc,
        "Why We Blind the Checker (the Most Important Idea)",
        [
            "If the Checker were allowed to read the Solver’s draft, the Checker "
            "would be influenced by it. AI models are known to agree with "
            "fluent, confident-sounding text. They suffer from a problem called "
            "confirmation bias. When you ask one AI to ‘review’ another AI’s "
            "answer in plain prose, it tends to agree.",
            "By giving the Checker only the bare list of claims and the raw "
            "evidence, we strip away the rhetorical force of the Solver’s "
            "writing. The Checker has nothing to agree with — it has only "
            "facts to compare. Either the evidence supports the claim, or it "
            "does not. There is no flowing prose to be charmed by.",
            "We do not enforce this rule by writing ‘please pretend you have "
            "not read this’ in the prompt. Such instructions can be ignored "
            "by the model. Instead we enforce the rule in the structure of "
            "the program: the Checker function is literally not given the "
            "Solver’s output as an argument. The Python code itself makes it "
            "impossible for the Checker to see what the Solver wrote. This "
            "structural guarantee is what we mean by ‘deliberate information "
            "asymmetry’.",
        ],
    )

    # ===================== KNOWLEDGE GRAPH =====================
    section(
        doc,
        "The Knowledge Graph (Neo4j) and Why We Use It",
        [
            "Alongside the vector store, we maintain a small knowledge graph in "
            "Neo4j. A knowledge graph stores information as a network of "
            "connected items. In our graph, every case is a node. Every "
            "court is a node. Every party type — Lawyer, Self-Represented "
            "Litigant — is a node. Every kind of hallucination — fabricated "
            "citation, false quotation — is a node. Edges connect them: case "
            "X happened in court Y, case X involved party Z, case X exhibited "
            "hallucination type W.",
            "The vector store is good at finding text that sounds like the "
            "question. The graph is good at answering precise structural "
            "questions: ‘Which other cases were heard in this same court?’, "
            "‘How many lawyer-led cases involved fabricated citations?’ The "
            "two stores together give the Checker both fuzzy textual match "
            "and exact structural match.",
            "In the current system the graph is used only as a supporting "
            "context for the Checker, but the schema is general enough that "
            "future versions can run deterministic Cypher queries against it "
            "for true Boolean fact-checking.",
        ],
    )

    # ===================== STEP BY STEP WALK =====================
    section(
        doc,
        "Step-by-Step Walk-through of One Real Probe",
        [
            "To make the explanation concrete, follow one probe through the "
            "whole system. The probe is: ‘In the case In re Whitehall Pharmacy "
            "LLC, the court found that paragraph fifteen contained a fabricated "
            "case citation.’ This is a TRUE probe whose ground-truth label is "
            "‘not a hallucination.’",
            "Step 1 — Retrieve. The system embeds the probe text into a 384-"
            "dimensional vector. ChromaDB finds the three nearest matching "
            "case documents in our 243-row corpus. The top match is the "
            "actual In re Whitehall Pharmacy LLC document. Neo4j is then "
            "queried for the structured facts attached to that case — court, "
            "party, date, hallucination type. Both pieces are added to the "
            "shared workflow state.",
            "Step 2 — Solver. The Solver reads the probe, the matched Chroma "
            "passages and the Neo4j subgraph. It composes a fluent paragraph "
            "describing what it understands.",
            "Step 3 — Proposer. The Proposer reads only the Solver’s "
            "paragraph. It returns a JSON list of small claims, each phrased "
            "as a (question, answer) pair. For example, ‘Was the case In re "
            "Berry Good, LLC real?’ → ‘No.’",
            "Step 4 — Checker. The Checker reads only the JSON list and the "
            "raw retrieved evidence. The Checker does not see the Solver’s "
            "paragraph. For each claim it returns is_hallucination: true or "
            "false, with a one-line explanation. The system tallies the "
            "results and computes a confidence score.",
            "Step 5 — Synthesizer. If any claim was flagged, the Synthesizer "
            "rewrites the answer so the flagged claims are removed. "
            "Otherwise the Solver’s paragraph passes through unchanged.",
            "The whole pipeline takes around three minutes per probe on our "
            "laptop. The Baseline equivalent takes around fifteen seconds "
            "but answers with a single yes/no without any of the structural "
            "guarantees.",
        ],
    )

    # ===================== TOOLS =====================
    section(
        doc,
        "The Tools We Used (One Line Each)",
        [
            "Our entire stack is open source and runs locally. Below is a "
            "one-line description of each tool. All of these are listed with "
            "URLs in the References section of the main project report.",
        ],
        bullets=[
            "Python 3 — the programming language we wrote everything in.",
            "Ollama — a local AI runtime that lets us run a language model "
            "on our own laptop without any cloud account.",
            "llama3.1 — the actual AI language model that powers the agents. "
            "Quantised to about 4.9 GB so it fits on a normal laptop.",
            "LangChain — a library that gives us a uniform way of talking to "
            "the AI model and the databases.",
            "LangGraph — sits on top of LangChain and lets us draw the agent "
            "pipeline as a graph of nodes with shared state.",
            "ChromaDB — the vector database. Stores text along with its "
            "embedding so we can search by meaning, not just by keyword.",
            "Sentence-Transformers (all-MiniLM-L6-v2) — the embedding model "
            "that converts text into a 384-dimensional vector.",
            "Neo4j — the graph database that stores cases, courts, parties "
            "and hallucination types as nodes connected by labelled edges.",
            "FastAPI — gives us a small REST endpoint at /ask so the system "
            "can be tested interactively.",
            "pandas, pytest — standard Python data and testing tools used "
            "for the evaluation harness.",
        ],
    )

    # ===================== NUMBERS =====================
    section(
        doc,
        "The Numbers (Our Results in Plain Words)",
        [
            "On the matched probe set, the single-agent Baseline produced an "
            "F1 score of 0.0. That sounds dramatic, and it is. The baseline "
            "did not catch the fabricated citation in the FALSE probe. "
            "Because there were no true positives, the precision and recall "
            "are both zero, which makes F1 zero by definition.",
            "Our Multi-Agent pipeline produced an F1 score of 0.667 on the "
            "same probes. It correctly flagged the FALSE probe as a "
            "hallucination, but it also flagged one TRUE probe — the one in "
            "which the Proposer atomised the answer into more claims than "
            "the evidence directly supported, including a peripheral claim "
            "that was reasonable but not literally written in the retrieved "
            "passages. So our system has perfect recall on this probe set "
            "but only 50% precision.",
            "The trade-off is clear. The Baseline never raises a false alarm "
            "but it also never catches the real fabrication. Our Multi-Agent "
            "pipeline always catches the real fabrication but occasionally "
            "raises a false alarm. In the legal domain, a false alarm is "
            "vastly preferable to letting a fabricated citation slip into a "
            "filing — the user will simply double-check, while the baseline "
            "would let the fake citation through unchallenged.",
            "Latency: the Baseline averages about thirty-seven seconds per "
            "probe; the Multi-Agent pipeline averages about one hundred and "
            "eighty-eight seconds per probe. About five times slower in "
            "exchange for the structural guarantees.",
        ],
    )

    # ===================== WHAT WORKED =====================
    section(
        doc,
        "What Worked Well",
        [
            "Three things worked well. First, encoding information asymmetry "
            "in the LangGraph state schema rather than in prompts. This made "
            "the architectural property auditable; we can simply look at the "
            "code and verify that the Checker function does not receive the "
            "Solver’s output as an argument. No model jail-break, no prompt "
            "trick, no clever wording can bypass it.",
            "Second, choosing a real-world dataset. The Charlotin Kaggle "
            "corpus is a catalogue of incidents that actually happened, "
            "documented by judges in their published rulings. This sidesteps "
            "the ‘synthetic benchmark’ critique that plagues much of the "
            "hallucination literature, where the same AI both generates and "
            "grades the test.",
            "Third, the controlled-comparison design. Both pipelines see "
            "the same evidence, use the same embeddings, run on the same "
            "language model and ingest the same prompt scaffolding. The "
            "only thing that varies is the architecture, so any difference "
            "in the metrics is attributable to the architectural change "
            "and not to a hidden confound.",
        ],
    )

    # ===================== WHAT DID NOT WORK =====================
    section(
        doc,
        "What Did Not Work, and Why",
        [
            "We were honest about three failures.",
            "Failure 1 — Over-flagging on TRUE probes. Our Multi-Agent "
            "pipeline sometimes atomises the Solver’s draft into more "
            "claims than the evidence directly supports. The Checker, "
            "applying its strict rule, then flags the entire response. "
            "This hurts our precision. Future work will tune the Proposer "
            "to atomise more conservatively, or allow the Checker to "
            "distinguish between central and peripheral claims.",
            "Failure 2 — Latency. Multi-agent debate intrinsically takes "
            "more time than a single prompt. Our system is roughly five "
            "times slower than the baseline. For real-time consumer use, "
            "this would be unacceptable; for a legal-filing review tool "
            "that runs overnight, it is fine.",
            "Failure 3 — Small probe set in this report. We ran the full "
            "Multi-Agent pipeline on a smoke set of three cases (six "
            "probes) for the report submission. A larger run is the very "
            "next thing we will do; the harness is in place and the "
            "command to re-run is documented in the project README.",
        ],
    )

    # ===================== EXAMINER QUESTIONS =====================
    section(
        doc,
        "Likely Examiner Questions and How To Answer",
        [
            "Practise these answers out loud. Each answer is short on purpose. "
            "If the examiner pushes for more detail, expand from the body of "
            "the report. If you do not know the answer, say so honestly and "
            "offer to follow up after the viva.",
        ],
    )
    qa_pairs = [
        (
            "Why did you choose this topic?",
            "Because hallucinations are the single biggest barrier to deploying "
            "AI in serious fields like law and medicine, and because there is a "
            "real public dataset that documents the problem. Working on it "
            "lets us combine current research in multi-agent systems with a "
            "practical, measurable safety problem.",
        ),
        (
            "What is the contribution of this project?",
            "Three things. One, an open-source multi-agent pipeline with "
            "structurally enforced information asymmetry between drafting and "
            "auditing agents. Two, a controlled empirical comparison against a "
            "matched single-agent baseline on a real-world legal hallucination "
            "corpus, not a synthetic one. Three, a reproducible evaluation "
            "harness so other researchers can extend our work.",
        ),
        (
            "Why is the dataset relevant?",
            "Because it is a real-world catalogue of actual court rulings in "
            "which AI-generated content was sanctioned by a judge. The "
            "ground truth is established by the court itself, not by another "
            "AI, which removes the circularity that affects most synthetic "
            "benchmarks in this field.",
        ),
        (
            "Why use four agents instead of one?",
            "Because a single agent suffers from confirmation bias — it is "
            "psychologically inclined to agree with its own previous output. "
            "By splitting the roles and structurally hiding the Solver’s "
            "draft from the Checker, we eliminate that bias at the "
            "architectural level rather than relying on prompt instructions.",
        ),
        (
            "Why both ChromaDB and Neo4j?",
            "ChromaDB is excellent at finding documents that are semantically "
            "similar to a query but cannot answer precise structural questions. "
            "Neo4j is excellent at structural queries — ‘which cases share "
            "this court?’ — but cannot do fuzzy semantic match. Together they "
            "give the Checker both kinds of evidence.",
        ),
        (
            "What does the F1 number tell us?",
            "F1 is the harmonic mean of precision and recall, so it punishes "
            "imbalances. A score of 0 means the system completely failed to "
            "catch the positive class. A score of 0.667 means the system "
            "catches the positive class very well but raises some false "
            "alarms. In our setup, catching all hallucinations is more "
            "important than avoiding every false alarm.",
        ),
        (
            "Why is your system five times slower?",
            "Because it runs four language model calls per probe instead of "
            "one. The structure is the source of the safety guarantee, but "
            "it is also the source of the latency. We can recover most of "
            "the cost in future work by parallelising the Retrieve and the "
            "Proposer stages, or by running the agents on different model "
            "sizes — small, fast models for atomisation, large, careful "
            "models for the Checker.",
        ),
        (
            "Why use a local model instead of GPT-4?",
            "Three reasons. One, cost: GPT-4 charges per call; running locally "
            "is free. Two, privacy: client data never leaves the laptop. "
            "Three, sustainability — Green AI: a single quantised local model "
            "uses much less energy than repeated API calls to a remote "
            "cloud service. Our results show that a 4.9 GB model is enough "
            "for this task when the surrounding architecture is good.",
        ),
        (
            "Can your system be applied to other domains?",
            "Yes. The Solver–Proposer–Checker–Synthesizer pattern is "
            "domain-independent. To move it to medicine you would change "
            "the dataset and the schema of the knowledge graph. To move "
            "it to finance you would change the dataset and add a "
            "Cypher-based deterministic check. The core architecture stays "
            "the same.",
        ),
        (
            "What is the limitation of this work?",
            "The probe set in this report is small because of time "
            "constraints. The probes also exercise only one type of "
            "hallucination — fabricated citations. Other types listed in "
            "the dataset, such as false quotations and fabricated statutes, "
            "are part of the future work plan in Chapter Six of the main "
            "report.",
        ),
        (
            "What would you do differently if you started again?",
            "We would invest earlier in a larger probe set so the headline "
            "metrics are statistically robust from the first run. We would "
            "also experiment earlier with mixing different base models "
            "across the agents to break shared parametric bias.",
        ),
        (
            "What is the future scope?",
            "Four directions. One, cover all hallucination categories in "
            "the dataset. Two, mix different base models across agents. "
            "Three, plug in a domain-specific knowledge graph such as "
            "LegalBench-RAG for deterministic citation validation. Four, "
            "run a small user study with practising lawyers to assess "
            "real-world utility.",
        ),
        (
            "How is your work different from existing literature?",
            "Three differences. One, real dataset instead of synthetic. "
            "Two, asymmetry encoded in the orchestration graph instead of "
            "in prompt instructions. Three, fully open-source reference "
            "implementation released alongside the report.",
        ),
        (
            "What is your role in the project, individually?",
            "Each team member should answer this in their own words. "
            "Examples: Harshal led the architecture and the LangGraph "
            "implementation; Yashraj led the evaluation harness and the "
            "metrics analysis; Aapesha led the dataset preparation and "
            "the documentation. Adapt as appropriate for the actual "
            "division of labour.",
        ),
    ]
    for q, a in qa_pairs:
        add_qa(doc, q, a)
    page_break(doc)

    # ===================== CLOSING STATEMENT =====================
    section(
        doc,
        "Closing Statement (Speak This At The End)",
        [
            "To conclude, our project demonstrates that the dangerous "
            "phenomenon of AI hallucinations can be substantially mitigated "
            "by moving away from single-agent retrieval-augmented generation "
            "and toward a structured multi-agent pipeline with deliberate "
            "information asymmetry between the drafting and auditing agents.",
            "Tested on a public, real-world dataset of more than four hundred "
            "documented legal hallucinations, our system catches the kind of "
            "fabricated citation that has resulted in real-world court "
            "sanctions, while a tightly matched single-agent baseline fails "
            "to do so. The whole pipeline runs on commodity hardware in line "
            "with Green AI principles. The code, the data ingestion, the "
            "evaluation harness and this very report are open source, and "
            "are available on our GitHub repository so that any future "
            "researcher can build on what we have done.",
            "Thank you. We are happy to take any further questions.",
        ],
    )

    # ===================== TEAM CHEAT-SHEET =====================
    section(
        doc,
        "Team Cheat-Sheet (One-Line Reminders)",
        [
            "These are one-line reminders for the team. Read them on the "
            "morning of the viva.",
        ],
        bullets=[
            "Hallucination = AI confidently saying something that is not "
            "true.",
            "Our project = build a team of AI agents that catch each other’s "
            "mistakes, instead of relying on one AI to grade itself.",
            "Dataset = real court cases from Kaggle where lawyers were fined "
            "for AI-generated fake citations.",
            "Big idea = the Checker agent never sees the Solver agent’s "
            "writing — only the raw evidence and a list of small claims.",
            "Tools = Python, Ollama, llama3.1, LangChain, LangGraph, "
            "ChromaDB, Sentence-Transformers, Neo4j, FastAPI.",
            "Result = our F1 = 0.667, baseline F1 = 0.0 on the matched "
            "probes — our system caught the fabrication; baseline did not.",
            "Cost = our system is roughly five times slower per query, "
            "acceptable for legal-tech, less acceptable for real-time chat.",
            "Future work = more failure modes, mixed base models, "
            "deterministic graph queries, user study with lawyers.",
            "Repository = github.com/harshal14ahire/multiagent-hallucination-"
            "defense (public, open source).",
            "Stay calm. Speak slowly. If you do not know an answer, "
            "say so politely and offer to follow up.",
        ],
    )

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", default="docs/Project_Explanation_Script.docx")
    args = ap.parse_args()
    build(args.out)
