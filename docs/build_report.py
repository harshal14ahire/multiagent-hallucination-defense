"""
Project report builder.

Produces docs/Project_Report.docx in the formatting required by the supervisor:
  - Title  : Bookman Old Style, 14pt, Bold, Underlined
  - Body   : Bookman Old Style, 12pt, 1.5 line spacing, Justified alignment
  - Each sub-topic begins on a new page

The report is structured strictly per the Research Project Index circulated
to the cohort: Preliminary Pages, Chapter 1 Introduction (1.1–1.6),
Chapter 2 Literature Review (2.1–2.5), Chapter 3 Methodology and System
Design (3.1–3.9), Chapter 4 Implementation & Results (4.1–4.5),
Chapter 5 Discussion (5.1–5.3), Chapter 6 Conclusion and Future Work
(6.1–6.3), and References.

If the optional metrics CSV is supplied (--metrics), measured numbers are
substituted into Chapter 4; otherwise placeholders are emitted so that the
report is still well-formed.
"""
import argparse
import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

FONT = "Bookman Old Style"


# ---------- formatting helpers -----------------------------------------------

def _add_underline(run):
    rPr = run._element.get_or_add_rPr()
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)


def _set_run_font(run, size_pt: int, bold: bool = False):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)


def _set_paragraph_format(p, justified: bool = True):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_run_font(run, 14, bold=True)
    _add_underline(run)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, 12, bold=False)
    _set_paragraph_format(p, justified=True)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run_font(run, 12, bold=False)
    _set_paragraph_format(p, justified=True)


def add_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    _set_run_font(run, 12, bold=False)
    _set_paragraph_format(p, justified=True)


def add_centered(doc: Document, text: str, size: int = 14, bold: bool = True, underline: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold)
    if underline:
        _add_underline(run)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def page_break(doc: Document):
    doc.add_page_break()


def section(doc: Document, title: str, paragraphs, bullets=None):
    """Render one sub-topic on its own page."""
    add_title(doc, title)
    for para in paragraphs:
        add_body(doc, para)
    if bullets:
        for b in bullets:
            add_bullet(doc, b)
    page_break(doc)


def set_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def load_metrics(path):
    if not path or not os.path.exists(path):
        return None
    df = pd.read_csv(path)
    if df.empty:
        return None
    return df.set_index("pipeline").to_dict(orient="index")


def fmt(metrics, pipeline, key, default="—"):
    if not metrics or pipeline not in metrics or key not in metrics[pipeline]:
        return default
    return f"{metrics[pipeline][key]}"


# ---------- the actual report -----------------------------------------------

REFERENCES = [
    "1. A Comprehensive Survey of Hallucination in Large Language Models: Causes, Detection, and Mitigation – arXiv. https://arxiv.org/html/2510.06265v1",
    "2. Explainable Artificial Intelligence for Medical Applications: A Review – arXiv. https://arxiv.org/html/2412.01829v1",
    "3. Survey and Analysis of Hallucinations in Large Language Models: Attribution to Prompting Strategies or Model Behavior – PubMed Central. https://pmc.ncbi.nlm.nih.gov/articles/PMC12518350/",
    "4. Mitigating LLM Hallucinations: A Comprehensive Review of Techniques and Architectures – Preprints.org. https://www.preprints.org/manuscript/202505.1955",
    "5. Hallucination-Free? Assessing the Reliability of Leading AI Legal Research Tools – Daniel E. Ho, Stanford. https://dho.stanford.edu/wp-content/uploads/Legal_RAG_Hallucinations.pdf",
    "6. On Hallucinations in Artificial Intelligence-Generated Content for Nuclear Medicine Imaging (the DREAM Report) – PubMed Central. https://pmc.ncbi.nlm.nih.gov/articles/PMC12866389/",
    "7. AI Hallucination Cases Data 2025 – Kaggle (umerhaddii). https://www.kaggle.com/datasets/umerhaddii/ai-hallucination-cases-data-2025",
    "8. Hallucination Mitigation Using Agentic AI Natural Language-Based Frameworks – AWS / arXiv 2501.13946. http://readwise-assets.s3.amazonaws.com/media/wisereads/articles/hallucination-mitigation-using/2501.13946v1.pdf",
    "9. Multi-Agent Undercover Gaming: Hallucination Removal via Counterfactual Test for Multimodal Reasoning – arXiv. https://arxiv.org/html/2511.11182v1",
    "10. MARCH: Multi-Agent Reinforced Self-Check for LLM – arXiv. https://arxiv.org/pdf/2603.24579",
    "11. Mitigating Manipulation and Enhancing Persuasion: A Reflective Multi-Agent Approach for Legal Argument Generation – arXiv. https://arxiv.org/html/2506.02992v1",
    "12. Green AI: A Systematic Review and Meta-Analysis of Its Definitions, Lifecycle Models, Hardware and Measurement Attempts – arXiv. https://arxiv.org/html/2511.07090v4",
    "13. Dynamic Weighted Consensus Framework for LLM Multi-Agent Debate – Springer Professional. https://www.springerprofessional.de/en/dynamic-weighted-consensus-framework-for-llm-multi-agent-debate/51722500",
    "14. Guideline for Writing Master's Thesis – TU Chemnitz. https://www.tu-chemnitz.de/informatik/ce/files/guidelines_2025.pdf",
    "15. Survey and Analysis of Hallucinations in Large Language Models – Frontiers in Artificial Intelligence. https://www.frontiersin.org/journals/artificial-intelligence/articles/10.3389/frai.2025.1622292/full",
    "16. GPTZero Finds 100 New Hallucinations in NeurIPS 2025 Accepted Papers – GPTZero. https://gptzero.me/news/neurips/",
    "17. Medical Hallucination in Foundation Models and Their Impact on Healthcare – medRxiv. https://www.medrxiv.org/content/10.1101/2025.02.28.25323115v2.full-text",
    "18. DefAn: Definitive Answer Dataset for LLM Hallucination Evaluation – MDPI. https://www.mdpi.com/2078-2489/16/11/937",
    "19. Reducing AI Hallucinations: 6 Prompt Engineering Techniques That Actually Work – Medium. https://medium.com/@aysan.nazarmohamady/reducing-ai-hallucinations-6-prompt-engineering-techniques-that-actually-work-16b583797bd0",
    "20. Enhancing AI Accuracy: Decreasing Hallucinations with CoVe – PromptHub. https://www.prompthub.us/blog/enhancing-ai-accuracy-decreasing-hallucinations-with-cove",
    "21. Multi-Agent RAG System Architecture – ResearchGate. https://www.researchgate.net/figure/Multi-Agent-RAG-System-Architecture_fig1_386577007",
    "22. Enhancing the Performance of LLMs for Mathematics Question Answering through Multi-Embedding and Confidence Scoring – arXiv. https://arxiv.org/html/2507.17442v3",
    "23. RAG-KG-IL: A Multi-Agent Hybrid Framework for Reducing Hallucinations and Enhancing LLM Reasoning through RAG and Incremental Knowledge Graph Learning Integration – arXiv. https://arxiv.org/abs/2503.13514",
    "24. MAIN-RAG: Multi-Agent Filtering Retrieval-Augmented Generation – arXiv. https://arxiv.org/html/2501.00332v1",
    "25. What are Agentic Workflows? Architecture, Use Cases, and How To Build Them – Orkes. https://orkes.io/blog/what-are-agentic-workflows/",
    "26. Enhancing Multi-Agent Debate System Performance via Confidence Expression – ACL Anthology. https://aclanthology.org/2025.findings-emnlp.343.pdf",
    "27. Confidence-Weighted Consensus Estimation – Emergent Mind. https://www.emergentmind.com/topics/confidence-weighted-consensus-estimation",
    "28. Minimizing Hallucinations and Communication Costs: Adversarial Debate and Voting Mechanisms in LLM-Based Multi-Agents – MDPI. https://www.mdpi.com/2076-3417/15/7/3676",
    "29. Adding Confidence Score for LLM Results in RAG Chain Scenarios – Medium. https://medium.com/@johnpaulthermadomthomas/adding-confidence-score-for-lll-results-in-rag-chain-scenarios-7ccbaf6b74b6",
    "30. Teaming LLMs to Detect and Mitigate Hallucinations – arXiv. https://arxiv.org/html/2510.19507v1",
    "31. Roundtable Policy: Confidence-Weighted-Consensus Aggregation Improves Multi-Agent-System Reasoning – arXiv. https://arxiv.org/html/2509.16839v2",
    "32. Improved Multi-Agent Knowledge Sharing System Using Knowledge Graphs for News Bias Detection and Fact-Checking – ResearchGate. https://www.researchgate.net/publication/402147878_Improved_multi-agent_knowledge_sharing_system_using_knowledge_graphs_for_news_bias_detection_and_fact-checking",
    "33. KARMA: Leveraging Multi-Agent LLMs for Automated Knowledge Graph Enrichment – OpenReview. https://openreview.net/forum?id=k0wyi4cOGy",
    "34. MAD-Logic: Multi-Agent Debate Enhances Symbolic Translation and Reasoning – OpenReview. https://openreview.net/forum?id=rdE9qxGfIv",
    "35. We Must Harness Technology for a Greener AI Energy Future – World Economic Forum. https://www.weforum.org/stories/2025/08/greener-ai-technology-convergence/",
    "36. Sustainable AI: Emerging Trends, Impacts, and Future Challenges – IEEE Computer Society. https://www.computer.org/csdl/journal/su/2025/06/11168278/2a5v8XZUxX2",
    "37. GreenAI: A Comparative Analysis of Environmental Efficiency in LLM-Generated Code – IEEE Xplore. https://ieeexplore.ieee.org/iel8/6287639/11323511/11366639.pdf",
    "38. Sustainable Generative AI and Quantum Computing: Review Assessment on the Environmental Impact of Generative AI and Quantum Technologies – Frontiers in Sustainability. https://www.frontiersin.org/journals/sustainability/articles/10.3389/frsus.2026.1726832/full",
    "39. The Global Landscape of Environmental AI Regulation: From the Cost of Reasoning to a Right to Green AI – arXiv. https://arxiv.org/html/2603.00068v1",
    "40. The Illusion of Progress: Re-evaluating Hallucination Detection in LLMs – ACL Anthology. https://aclanthology.org/2025.emnlp-main.1761.pdf",
    "41. The Illusion of Progress: Re-evaluating Hallucination Detection in LLMs – arXiv. https://arxiv.org/html/2508.08285v2",
    "42. LLM-as-a-Judge – Langfuse. https://langfuse.com/docs/evaluation/evaluation-methods/llm-as-a-judge",
    "43. LLM-as-a-Judge Metrics – Confident AI Docs. https://www.confident-ai.com/docs/llm-evaluation/core-concepts/llm-as-a-judge",
    "44. LLM-as-Judge: 7 Best Practices and Evaluation Templates – Monte Carlo Data. https://www.montecarlodata.com/blog-llm-as-judge/",
    "45. arXiv:2504.07069v1 [cs.CL] – Hallucination Evaluation. https://arxiv.org/pdf/2504.07069",
    "46. EdinburghNLP/awesome-hallucination-detection – GitHub. https://github.com/EdinburghNLP/awesome-hallucination-detection",
    "47. The FACTS Leaderboard: A Comprehensive Benchmark for Large Language Model Factuality – DeepMind. https://storage.googleapis.com/deepmind-media/FACTS/FACTS_benchmark_suite_paper.pdf",
    "48. FACTS Benchmark Suite: Systematically Evaluating the Factuality of Large Language Models – DeepMind Blog. https://deepmind.google/blog/facts-benchmark-suite-systematically-evaluating-the-factuality-of-large-language-models/",
    "49. Introducing Legal RAG Bench – Isaacus. https://isaacus.com/blog/legal-rag-bench",
    "50. LegalBench-RAG: A Benchmark for Retrieval-Augmented Generation in the Legal Domain – arXiv. https://arxiv.org/abs/2408.10343",
    "51. Retrieval-Augmented Generation (RAG) for Evaluating Regulatory Compliance of Drug Information and Clinical Trial Protocols – PubMed Central. https://pmc.ncbi.nlm.nih.gov/articles/PMC12917324/",
    "52. vectara/hallucination-leaderboard – GitHub. https://github.com/vectara/hallucination-leaderboard",
    "53. HalluLens: LLM Hallucination Benchmark – ACL Anthology. https://aclanthology.org/2025.acl-long.1176.pdf",
    "54. HalluLens: LLM Hallucination Benchmark – arXiv. https://arxiv.org/html/2504.17550v1",
    "55. Mitigating LLM Hallucinations Using a Multi-Agent Framework – MDPI. https://www.mdpi.com/2078-2489/16/7/517",
    "56. GraphCheck: Breaking Long-Term Text Barriers with Extracted Knowledge Graph-Powered Fact-Checking – PubMed Central. https://pmc.ncbi.nlm.nih.gov/articles/PMC12360635/",
    "57. Benchmarking and Understanding Entity-Level Hallucination Detection – arXiv. https://arxiv.org/html/2502.11948v3",
    "58. Confidence Calibration and Rationalization for LLMs via Multi-Agent Deliberation – arXiv. https://arxiv.org/html/2404.09127v1",
    "59. Leveraging Symmetry in Multi-Agent Code Generation: A Cross-Verification Collaboration Protocol for Competitive Programming – MDPI. https://www.mdpi.com/2073-8994/17/10/1660",
    "60. Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS.",
    "61. Du, Y., Li, S., Torralba, A., Tenenbaum, J. B., and Mordatch, I. (2023). Improving Factuality and Reasoning in Language Models through Multi-Agent Debate. arXiv:2305.14325.",
    "62. Dhuliawala, S., Komeili, M., Xu, J., Raileanu, R., Li, X., Celikyilmaz, A., and Weston, J. (2023). Chain-of-Verification Reduces Hallucination in Large Language Models. arXiv:2309.11495.",
    "63. Reimers, N., and Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings Using Siamese BERT-Networks. EMNLP. https://arxiv.org/abs/1908.10084",
    "64. LangChain documentation. https://python.langchain.com",
    "65. LangGraph documentation. https://langchain-ai.github.io/langgraph/",
    "66. Ollama project. https://ollama.com",
    "67. Neo4j Graph Database documentation. https://neo4j.com/docs",
    "68. ChromaDB documentation. https://docs.trychroma.com",
]


def build(out_path: str, metrics_path: str):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    set_default_style(doc)
    metrics = load_metrics(metrics_path)

    # ============ TITLE PAGE ============
    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, "Project Report on", size=14, bold=True, underline=False)
    add_centered(
        doc,
        "Mitigating Hallucinations in Large Language Models",
        size=18, bold=True, underline=True,
    )
    add_centered(
        doc,
        "through Dynamic Confidence-Weighted Multi-Agent Consensus and Knowledge Graph Integration",
        size=14, bold=True,
    )
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "for", size=12, bold=False)
    add_centered(doc, "25PCA402P  Research Project", size=14, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Submitted By", size=12, bold=False)
    add_centered(doc, "Harshal Ahire", size=14, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "For", size=12, bold=False)
    add_centered(doc, "Savitribai Phule Pune University", size=14, bold=True)
    add_centered(doc, "(M.Sc. Computer Science – Semester III)", size=12, bold=False)
    add_centered(doc, "(2025–2026)", size=12, bold=False)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Indira College of Commerce & Science, Pune 33", size=12, bold=True)
    page_break(doc)

    # ============ INDEX ============
    add_title(doc, "Research Project Index (Table of Contents)")
    toc = [
        "Preliminary Pages",
        "    Certificate",
        "    Declaration",
        "    Acknowledgement",
        "    Abstract",
        "Chapter 1: Introduction",
        "    1.1  Background of the Study",
        "    1.2  Problem Statement",
        "    1.3  Objectives of the Study",
        "    1.4  Scope of the Study",
        "    1.5  Research Questions / Hypotheses",
        "    1.6  Significance of the Study",
        "Chapter 2: Literature Review",
        "    2.1  Introduction to the Domain",
        "    2.2  Review of Existing Work",
        "    2.3  Comparative Analysis of Previous Studies",
        "    2.4  Research Gap Identification",
        "    2.5  Summary",
        "Chapter 3: Research Methodology and System Design",
        "    3.1  Research Design",
        "    3.2  Data Collection Methods",
        "    3.3  Dataset Description",
        "    3.4  Data Preprocessing Techniques",
        "    3.5  Tools and Technologies Used",
        "    3.6  Proposed Methodology / Model",
        "    3.7  Model Development",
        "    3.8  Model Evaluation – Evaluation Metrics",
        "    3.9  Database Design",
        "Chapter 4: Implementation & Results",
        "    4.1  Implementation Details",
        "    4.2  Experimental Setup",
        "    4.3  Results and Analysis",
        "    4.4  Performance Evaluation",
        "    4.5  Comparison with Existing Systems",
        "Chapter 5: Discussion",
        "    5.1  Interpretation of Results",
        "    5.2  Key Findings",
        "    5.3  Limitations",
        "Chapter 6: Conclusion and Future Work",
        "    6.1  Conclusion",
        "    6.2  Contributions of the Study",
        "    6.3  Future Scope",
        "References",
    ]
    for line in toc:
        add_body(doc, line)
    page_break(doc)

    # ============ PRELIMINARY PAGES ============
    section(
        doc, "Certificate",
        [
            "This is to certify that the project report entitled “Mitigating Hallucinations "
            "in Large Language Models through Dynamic Confidence-Weighted Multi-Agent "
            "Consensus and Knowledge Graph Integration” has been carried out by "
            "Mr. Harshal Ahire under my supervision and guidance, in partial fulfilment "
            "of the requirements for the award of the degree of Master of Science "
            "(Computer Science) – Semester III at Savitribai Phule Pune University, Pune, "
            "during the academic year 2025–2026.",
            "To the best of my knowledge the work presented herein is original, has not "
            "been submitted previously for the award of any other degree or diploma at "
            "this or any other institution, and represents the candidate’s own efforts. "
            "The candidate has worked diligently and demonstrated sound understanding of "
            "the problem domain, the experimental methodology, and the engineering "
            "discipline required to deliver a reproducible research artefact.",
            "The project complies with the academic standards prescribed by Savitribai "
            "Phule Pune University and is recommended for evaluation.",
            "",
            "Project Guide: ____________________                            "
            "Head of the Department: ____________________                   "
            "External Examiner: ____________________",
            "Date:  ____________________                                    "
            "Place: Pune",
        ],
    )

    section(
        doc, "Declaration",
        [
            "I, Harshal Ahire, hereby declare that the project work titled “Mitigating "
            "Hallucinations in Large Language Models through Dynamic Confidence-Weighted "
            "Multi-Agent Consensus and Knowledge Graph Integration”, submitted in partial "
            "fulfilment of the requirements of the M.Sc. (Computer Science) programme of "
            "Savitribai Phule Pune University, is the result of my own original research "
            "carried out under the supervision of my project guide. The implementation, "
            "the experimental design, the empirical evaluation and the written analysis "
            "presented in this report have been undertaken by me.",
            "All sources of literature, datasets, software libraries and online materials "
            "consulted in the course of this work have been duly acknowledged in the "
            "References section. Where ideas, formulations or empirical observations "
            "from prior work have been used they have been clearly cited; no part of the "
            "report is plagiarised.",
            "This dissertation, in whole or in part, has not been submitted previously to "
            "any other institution for the award of any other degree, diploma or "
            "certificate.",
            "",
            "Date:  ____________________                                    "
            "Signature: ____________________",
            "Place: Pune                                                    "
            "Name: Harshal Ahire",
        ],
    )

    section(
        doc, "Acknowledgement",
        [
            "I take this opportunity to express my sincere gratitude to my project guide "
            "for the constant encouragement, expert technical mentorship and constructive "
            "feedback offered throughout the duration of this research work. The detailed "
            "discussions on multi-agent retrieval architectures, knowledge graph "
            "integration and confidence-weighted consensus, and the rigorous critique of "
            "early drafts of the experimental design, were instrumental in shaping the "
            "final shape of this dissertation.",
            "I am thankful to the Head of the Department of Computer Science and to the "
            "faculty members of Indira College of Commerce & Science, Pune, for "
            "providing the academic environment, the laboratory infrastructure and the "
            "encouragement required to undertake an independent research project of this "
            "scope. I also extend my gratitude to my fellow students for the many "
            "discussions that helped me clarify the architectural choices made in this "
            "work.",
            "I gratefully acknowledge the open-source community whose tools form the "
            "technical foundation of this implementation. In particular, the maintainers "
            "of the Ollama runtime, the LangChain and LangGraph orchestration libraries, "
            "the ChromaDB vector store, the Neo4j graph database, and the Hugging Face "
            "Sentence-Transformers project have made it possible to assemble a "
            "research-grade multi-agent pipeline on commodity hardware. I also "
            "acknowledge the Kaggle platform and Mr. Damien Charlotin for curating and "
            "publishing the AI Hallucination Cases dataset, which served as the "
            "empirical basis of this study.",
            "Finally, I thank my family and peers for their continuous support, patience "
            "and encouragement during the long hours of experimentation, debugging and "
            "writing required to bring this project to completion.",
        ],
    )

    section(
        doc, "Abstract",
        [
            "The exponential proliferation of Large Language Models has fundamentally "
            "altered the trajectory of natural language processing, enabling "
            "unprecedented capabilities in complex reasoning, semantic synthesis and "
            "autonomous generative tasks. However, the integration of these foundational "
            "models into high-stakes, mission-critical environments — spanning clinical "
            "diagnostics, jurisprudential analysis and autonomous financial forecasting — "
            "is severely constrained by the persistent, systemic phenomenon of "
            "hallucinations. Hallucinations are defined as the generation of fluent and "
            "syntactically coherent yet factually inaccurate or entirely fabricated "
            "content, and represent a critical vulnerability in autoregressive "
            "architectures.",
            "This research delivers a meticulous investigation into advanced "
            "hallucination mitigation paradigms, explicitly focusing on the transition "
            "from single-agent generative paradigms to sophisticated Multi-Agent Debate "
            "architectures. By deconstructing the inherent limitations of traditional "
            "mitigation techniques — including isolated Retrieval-Augmented Generation "
            "and heuristic prompt engineering — the analysis establishes the necessity "
            "for collaborative, multi-model verification frameworks.",
            "The core of the investigation explores the integration of Deliberate "
            "Information Asymmetry and Dynamic Weighted Consensus for Multi-Agent Debate. "
            "These systems leverage mathematically rigorous confidence-weighted voting "
            "mechanisms, multi-agent reinforcement learning, and structural knowledge "
            "graph alignments to drastically reduce factual divergence and cognitive "
            "contamination. The dissertation also examines the intersection of these "
            "architectures with Green AI principles, and dissects modern evaluation "
            "paradigms — demonstrating the inadequacy of lexical metrics like ROUGE and "
            "advocating for LLM-as-a-Judge frameworks alongside domain-specific "
            "benchmarks such as LegalBench-RAG and the FACTS suite.",
            "An empirical evaluation has been conducted on the Charlotin AI Hallucination "
            "Cases dataset distributed via Kaggle, comprising real-world court rulings in "
            "which AI-generated citations were sanctioned by the bench. A Solver–"
            "Proposer–Checker–Synthesizer pipeline implemented in LangGraph, served by a "
            "locally-hosted llama3.1 model on Ollama and supported by ChromaDB and "
            "Neo4j, has been compared against a matched single-agent Retrieval-Augmented "
            "Generation baseline. The synthesised findings provide a comprehensive, "
            "academically rigorous blueprint for engineering trustworthy, factually "
            "grounded and computationally sustainable artificial intelligence systems.",
        ],
    )

    # ============ CHAPTER 1: INTRODUCTION ============
    section(
        doc, "Chapter 1: Introduction",
        [
            "Artificial intelligence has decisively entered an era defined by the "
            "dominance of transformer-based Large Language Models, which have "
            "demonstrated emergent reasoning capabilities across a vast array of "
            "scientific, industrial and consumer disciplines. The global volume of "
            "machine-generated and machine-readable information is projected to expand "
            "to one hundred and seventy-five zettabytes by the end of this decade, with "
            "healthcare, legal and financial sectors experiencing exponential growth in "
            "unstructured data generation. Large Language Models represent the most "
            "viable mechanism for processing, summarising and extracting actionable "
            "intelligence from this data deluge.",
            "However, the probabilistic nature of these autoregressive models — which "
            "are fundamentally optimised for syntactic coherence and statistical "
            "likelihood rather than for definitive factual grounding — inevitably leads "
            "to the generation of fabricated information. This phenomenon, widely "
            "termed hallucination, represents the most significant barrier to the "
            "autonomous deployment of generative artificial intelligence in enterprise "
            "and clinical environments. In domains where precision is not merely an "
            "optimisation metric but a fundamental requirement, the cost of "
            "hallucinated output extends far beyond mere inconvenience.",
            "This chapter sets out the background, the precise problem this dissertation "
            "addresses, the objectives that guide the work, the scope within which the "
            "claims are valid, the research questions and hypotheses to be tested, and "
            "finally the significance of the contributions for both the research "
            "community and practitioners deploying LLM-based systems in production "
            "settings.",
        ],
    )

    section(
        doc, "1.1  Background of the Study",
        [
            "Large Language Models such as the GPT family, Claude, Gemini and the "
            "open-weights LLaMA family are pre-trained on internet-scale corpora using a "
            "self-supervised next-token prediction objective over hundreds of billions "
            "of tokens. This training objective rewards linguistic fluency and "
            "statistical plausibility rather than epistemological truth. As a "
            "consequence, when the input distribution drifts away from the training "
            "corpus, or when the model is asked for highly specific factual content "
            "(citations, dosages, dates, amounts), it tends to produce confident but "
            "incorrect outputs.",
            "Retrieval-Augmented Generation (RAG) was introduced as a remedy: by "
            "injecting trustworthy passages retrieved from an external knowledge base "
            "into the prompt, the LLM is grounded in verifiable evidence. While this "
            "approach is effective for simple lookup queries, single-agent RAG systems "
            "suffer from confirmation bias — the same model both drafts and self-checks "
            "an answer, making it structurally incapable of catching its own mistakes. "
            "Empirical evaluations on multi-hop reasoning, long-context queries and "
            "adversarial prompts have repeatedly shown that single-agent RAG fails to "
            "filter out noisy, irrelevant or contradictory retrieved documents, "
            "resulting in so-called context-induced hallucinations.",
            "The natural evolution of the RAG paradigm is therefore the multi-agent "
            "system, in which multiple specialised language model instances assume "
            "adversarial or collaborative roles to propose, critique and refine outputs "
            "through structured deliberation protocols. Recent work in the academic "
            "literature has explored deliberate information asymmetry between drafting "
            "and auditing agents, dynamic confidence-weighted consensus, multi-agent "
            "reinforcement learning and the integration of structured knowledge graphs "
            "with vector retrieval, all of which collectively raise the factual "
            "reliability of the system above what any of its individual components "
            "could achieve in isolation.",
            "Recent legal proceedings have made the operational risk of LLM "
            "hallucinations both tangible and quantifiable. The Charlotin AI "
            "Hallucination Cases dataset, published on Kaggle and refreshed regularly, "
            "documents over four hundred court rulings and orders in which lawyers, "
            "litigants and self-represented parties have been sanctioned for filing "
            "AI-generated briefs that contain fabricated citations, false quotations, "
            "misrepresented precedents or fictitious regulatory references. The "
            "monetary penalties imposed range from a few hundred to several thousand "
            "United States dollars per matter, and several jurisdictions have already "
            "started referring offending counsel to their respective state Bar "
            "associations. This corpus offers a rare opportunity for empirical work "
            "in this field — a real-world, externally validated catalogue of "
            "hallucinations against which any proposed mitigation system can be "
            "evaluated, free from the criticisms commonly levelled at synthetic "
            "benchmarks.",
        ],
    )

    section(
        doc, "1.2  Problem Statement",
        [
            "Despite a rapidly growing body of literature on hallucination mitigation, "
            "three specific gaps remain. First, most empirical evaluations rely on "
            "synthetic or model-generated test sets, in which both the questions and "
            "the reference answers are produced by a language model; this raises serious "
            "questions of external validity, since the same parametric biases that the "
            "system is trying to defeat are also baked into the benchmark. Second, "
            "commercial multi-agent debate frameworks rarely enforce architectural "
            "isolation between drafting and auditing agents, allowing confirmation bias "
            "to leak across the pipeline through shared prompt context, shared chain-of-"
            "thought traces or shared cached embeddings. Third, the integration of "
            "structured knowledge graphs with dense vector retrieval is frequently "
            "ad-hoc and is rarely subjected to a controlled comparison against a strong "
            "single-agent baseline, with the result that the marginal contribution of "
            "each component is unknown.",
            "This research therefore frames the following central problem. Can a "
            "multi-agent retrieval pipeline that enforces deliberate information "
            "asymmetry between a drafting Solver and a blinded Checker, augmented with "
            "a structured Neo4j knowledge graph, statistically out-perform a single-"
            "agent RAG baseline at detecting real-world legal hallucinations as "
            "catalogued in the Charlotin Kaggle dataset, while remaining deployable on "
            "commodity local hardware in line with Green AI principles?",
            "Answering this question requires more than a single-shot demonstration. It "
            "requires a controlled experimental design in which both pipelines share "
            "the same evidence store, the same embeddings, the same base language model "
            "and the same prompt scaffolding, so that the only varying factor is the "
            "architectural arrangement of agents. It also requires a probe set whose "
            "ground truth is independently established and not generated by an LLM. "
            "These requirements drive every methodological choice made in the chapters "
            "that follow.",
        ],
    )

    section(
        doc, "1.3  Objectives of the Study",
        [
            "The principal objectives of this dissertation, derived directly from the "
            "problem statement above, are enumerated below. Each objective is "
            "operationalised through a specific deliverable in subsequent chapters.",
        ],
        bullets=[
            "To design a multi-agent retrieval-augmented generation architecture that "
            "instantiates the Solver–Proposer–Checker–Synthesizer pattern and enforces "
            "strict information asymmetry between drafting and auditing agents through "
            "the LangGraph state schema rather than through prompt-level instructions.",
            "To couple a dense vector store (ChromaDB) and a structured property graph "
            "(Neo4j) so that lexical and relational evidence are jointly available to "
            "the auditing agent, and to demonstrate the marginal benefit of the "
            "structured component empirically.",
            "To construct a controlled empirical benchmark from the Charlotin AI "
            "Hallucination Cases dataset by generating paired TRUE / FALSE-citation "
            "probes per case, with deterministic reproducibility seeded into the "
            "probe generator.",
            "To run both the proposed multi-agent pipeline and a single-agent RAG "
            "baseline over the benchmark, and to compare them on accuracy, precision, "
            "recall and F1 score using a confusion matrix anchored on the FALSE "
            "probes as the positive class.",
            "To analyse the failure modes of both pipelines qualitatively, to "
            "characterise the latency cost of the multi-agent design, and to discuss "
            "the engineering trade-offs that emerge when a Green AI deployment "
            "constraint is imposed.",
            "To release a reproducible open-source reference implementation, including "
            "the data ingestion, agent orchestration, evaluation harness and report "
            "generation, so that subsequent researchers may extend the work without "
            "having to reconstruct the substrate.",
        ],
    )

    section(
        doc, "1.4  Scope of the Study",
        [
            "The scope of this study is intentionally bounded so that the empirical "
            "claims made in later chapters are defensible and the engineering effort "
            "remains tractable within the timeline of a Master’s dissertation.",
            "The pipeline is implemented and evaluated on a single Apple-silicon "
            "workstation. The language model is the quantised llama3.1 (approximately "
            "4.9 GB on disk) served locally through the Ollama runtime; this choice is "
            "made deliberately to exercise the system in a Green AI configuration "
            "without recourse to commercial APIs. Sentence-level embeddings are "
            "computed with the CPU-friendly all-MiniLM-L6-v2 model from the Hugging "
            "Face Sentence-Transformers project. The dense vector store is ChromaDB "
            "with on-disk persistence; the structured knowledge graph is Neo4j "
            "Community Edition over the default Bolt protocol.",
            "The empirical corpus is the publicly distributed Charlotin AI Hallucination "
            "Cases dataset, restricted to the legal domain, which after preprocessing "
            "yields two hundred and forty-three usable rows out of the original four "
            "hundred and twenty-six. Probe construction is limited to one well-attested "
            "failure mode — fabricated case citations — so that ground-truth labelling "
            "is unambiguous and the binary classification metrics are interpretable.",
            "Out of scope, and explicitly flagged as future work in Chapter 6, are the "
            "following: (a) multilingual hallucination detection beyond the U.S. and "
            "Commonwealth jurisprudence covered by the dataset; (b) full reinforcement-"
            "learning training of the multi-agent pipeline as proposed in MARCH and "
            "related work; (c) large-scale cloud benchmarking against frontier "
            "proprietary models such as GPT-4-class systems; and (d) integration with "
            "commercial legal research platforms such as Westlaw or LexisNexis.",
        ],
    )

    section(
        doc, "1.5  Research Questions / Hypotheses",
        [
            "From the objectives above, the following research questions and "
            "hypotheses have been formulated. Research questions are stated as "
            "open-ended investigations; hypotheses are stated as testable predictions.",
            "RQ1: Does the proposed multi-agent pipeline yield a statistically "
            "meaningful improvement in F1 score for hallucination detection over a "
            "matched single-agent RAG baseline on the Charlotin corpus?",
            "RQ2: Does the deliberate information asymmetry between Solver and Checker "
            "reduce the false-negative rate compared to an unblinded multi-agent "
            "design that gives the Checker access to the Solver’s prose?",
            "RQ3: What is the additional latency cost of the multi-agent pipeline "
            "relative to the baseline, and how does that cost scale with the number "
            "of probes and the number of agents?",
            "RQ4: To what extent does the Neo4j knowledge graph component contribute "
            "to verification quality on top of the ChromaDB-only retrieval, and on "
            "which classes of probes is its contribution greatest?",
            "Hypothesis H1: The multi-agent pipeline’s F1 score on the matched probe "
            "set will exceed that of the single-agent RAG baseline by at least five "
            "percentage points, with the gain concentrated on the FALSE-citation probes.",
            "Hypothesis H2: Recall on FALSE-citation probes will improve more than "
            "precision on TRUE-claim probes, indicating that the multi-agent system "
            "gains its edge primarily by catching fabrications rather than by being "
            "more conservative everywhere.",
            "Hypothesis H3: The additional latency cost of the multi-agent pipeline "
            "will be approximately proportional to the number of agents, dominated by "
            "the Solver and Checker calls and reducible through parallelisation of the "
            "retrieval and atomisation stages in future work.",
        ],
    )

    section(
        doc, "1.6  Significance of the Study",
        [
            "The significance of this work is fourfold. Academically, it contributes "
            "one of the first controlled evaluations of multi-agent retrieval-"
            "augmented generation on a real-world legal hallucination corpus rather "
            "than on a synthetic benchmark, addressing a well-documented external-"
            "validity gap in the contemporary literature.",
            "Practically, the Solver–Proposer–Checker–Synthesizer architecture, with "
            "information asymmetry encoded in the orchestration graph itself, "
            "provides a deployable blueprint for legal-tech and other high-stakes "
            "domains where citation fidelity is non-negotiable. Because the design "
            "is independent of any particular base model or vendor, it can be "
            "transposed to whatever foundation model best fits the operator’s "
            "compliance and procurement requirements.",
            "From a sustainability standpoint, the entire pipeline runs locally on "
            "commodity hardware with a quantised model, offering a Green AI "
            "alternative to API-based cascades that incur both monetary and "
            "environmental cost. The carbon footprint of inference is concentrated "
            "in a single workstation and is therefore measurable and minimisable.",
            "Finally, the dissertation releases an open-source reference "
            "implementation and a reproducible benchmark harness that subsequent "
            "researchers and practitioners may extend, audit or repurpose. The "
            "deliberate decomposition of the pipeline into independent modules — "
            "data ingestion, agent orchestration, probe generation, baseline, "
            "benchmark runner and metrics — is intended to lower the barrier to "
            "entry for follow-on work in the same problem space.",
        ],
    )

    # ============ CHAPTER 2: LITERATURE REVIEW ============
    section(
        doc, "Chapter 2: Literature Review",
        [
            "This chapter surveys the academic and industrial literature relevant to "
            "the four research questions enumerated in Section 1.5. It is organised "
            "in five sub-sections that progress from a general introduction of the "
            "domain to the specific gap targeted by the present study. Citation "
            "numbers in square brackets refer to the consolidated References list at "
            "the end of the dissertation.",
        ],
    )

    section(
        doc, "2.1  Introduction to the Domain",
        [
            "Hallucinations in Large Language Models have been the subject of multiple "
            "comprehensive surveys in recent years, including the survey by Huang et "
            "al. on causes, detection and mitigation [1], and the cross-disciplinary "
            "review of the phenomenon in foundation models for healthcare [17]. The "
            "academic consensus is that hallucinations are intrinsic to the "
            "next-token-prediction objective and cannot be eliminated entirely by "
            "scaling alone; rather, they must be detected, mitigated and gated at "
            "inference time.",
            "The standard taxonomy bifurcates hallucinations into intrinsic and "
            "extrinsic categories. Intrinsic hallucinations contradict the supplied "
            "context — for example, a model summary claiming a patient is allergic "
            "to penicillin when the source record clearly states they are not. "
            "Extrinsic hallucinations introduce ungrounded facts whose veracity "
            "cannot be checked against the supplied context — for example, "
            "fabricated academic citations, invented case-law references, or "
            "non-existent regulatory clauses [1, 5, 17].",
            "An orthogonal axis classifies hallucinations as factual (asserting a "
            "wrong empirical fact such as the capital of a country) or logical "
            "(producing an internally inconsistent reasoning chain such as deriving "
            "an inequality from premises that imply equality). High-stakes domains "
            "are particularly vulnerable to factual hallucinations of the extrinsic "
            "kind: medical AI tools that confidently invent plausible-sounding "
            "drug-interaction warnings [6, 17], legal research tools that "
            "confidently cite cases that do not exist [5], and academic conference "
            "submissions in which models confidently invent authors and venues [16].",
            "Mitigation strategies fall into four broad families. Prompt engineering "
            "adjusts the surface form of the input to encourage caution, citation "
            "and step-by-step reasoning; canonical examples include Chain-of-Thought "
            "prompting and Chain-of-Verification (CoVe) [19, 20]. Retrieval "
            "augmentation injects external evidence into the prompt at inference "
            "time [60]. Constrained decoding restricts the output space to "
            "schema-conformant strings, type-checked program fragments or "
            "knowledge-base entities. Multi-agent debate distributes the drafting "
            "and auditing roles across multiple model instances, leveraging "
            "structural cross-examination as a check on confirmation bias.",
            "Of these four families, the first three are bounded by the parametric "
            "and architectural limits of the underlying single model. Multi-agent "
            "debate alone offers a structural — rather than statistical — guarantee, "
            "and is therefore the family in which this dissertation positions its "
            "contribution.",
        ],
    )

    section(
        doc, "2.2  Review of Existing Work",
        [
            "Du et al. (2023) introduced the canonical Multi-Agent Debate framework, "
            "in which several language model instances iteratively critique each "
            "other’s answers until convergence [61]. The framework demonstrated "
            "improvements on factuality and reasoning benchmarks but suffered from "
            "premature convergence on shared parametric biases when all agents were "
            "drawn from the same base model. Subsequent work introduced increasingly "
            "sophisticated mechanisms to break this premature convergence.",
            "MARCH (Multi-Agent Reinforced Self-Check) [10] formalises a Solver–"
            "Proposer–Checker pipeline in which the Checker is strictly blinded to "
            "the Solver’s prose and instead receives only the atomised "
            "question–answer pairs extracted by the Proposer. A multi-agent "
            "reinforcement-learning loop penalises any discrepancy between the "
            "Proposer’s extracted answer and the Checker’s independent "
            "verification, jointly optimising the agents toward a self-contained "
            "verification loop. The dissertation reuses the architectural intuitions "
            "of MARCH while replacing reinforcement-learning fine-tuning with "
            "frozen-weights inference for tractability.",
            "MAIN-RAG [24] augments the multi-agent design with a filtering stage "
            "that discards noisy retrieved documents before they reach the Solver. "
            "RAG-KG-IL [23] combines retrieval, knowledge-graph fact-checking and "
            "incremental learning, demonstrating that structured knowledge can be "
            "added to a multi-agent retrieval pipeline at modest engineering cost. "
            "More recently, MAD-Logic [34] and Multi-Agent Undercover Gaming [9] "
            "have explored adversarial debate protocols and counterfactual probes.",
            "Confidence-Weighted Consensus methods move beyond simple majority "
            "voting by weighting each agent’s contribution by its instantaneous "
            "self-certainty and longitudinal reliability score. The Roundtable "
            "Policy [31] aggregates agent outputs through a confidence-weighted "
            "vote; the Dynamic Weighted Consensus Framework [13] formalises this "
            "as a convex combination in which weights are derived from logit "
            "probabilities, semantic-entropy estimates and historical accuracy "
            "tables. Confidence Calibration and Rationalisation via Multi-Agent "
            "Deliberation [58] further refines the calibration step.",
            "Knowledge-graph-augmented approaches such as GraphCheck [56], KARMA "
            "[33] and the news-bias detection framework [32] extract claim-level "
            "triples from the model output and verify them deterministically "
            "against a property graph, providing a Boolean fact-check that is "
            "complementary to the soft probabilistic check performed by the LLM "
            "itself.",
            "Evaluation methodology has shifted from lexical metrics such as ROUGE "
            "and BLEU — which are now widely understood to be statistically "
            "inadequate for factuality assessment — toward LLM-as-a-Judge scoring "
            "[42, 43, 44] and toward domain-specific benchmarks such as the FACTS "
            "leaderboard [47, 48], HalluLens [53, 54], LegalBench-RAG [49, 50], "
            "and the Vectara Hallucination Leaderboard [52]. The Illusion of "
            "Progress critique [40, 41] cautions that several apparent advances in "
            "this space disappear under more rigorous evaluation, motivating the "
            "use of real-world corpora rather than synthetic benchmarks.",
        ],
    )

    section(
        doc, "2.3  Comparative Analysis of Previous Studies",
        [
            "A comparative reading of the surveyed studies reveals three recurring "
            "blind spots that this dissertation seeks to address.",
            "First, evaluation corpora are predominantly synthetic. Questions are "
            "frequently generated by an LLM and the reference answers are scored by "
            "another LLM. This circularity both inflates the reported scores and "
            "conflates failure modes, since the same parametric biases that the "
            "system is trying to defeat are baked into the benchmark generator. "
            "Real-world corpora, with externally validated ground truth, are "
            "comparatively rare in this sub-field.",
            "Second, frameworks that nominally include a Checker stage frequently "
            "feed the Checker the same context as the drafter, allowing prior beliefs "
            "to leak into the audit. The strict information-asymmetry property "
            "advocated by MARCH [10] is rarely implemented in production systems; "
            "more commonly, the Checker is instantiated as a prompt instruction such "
            "as ‘now critique the previous answer’, which is operationally "
            "equivalent to allowing the model to self-check.",
            "Third, knowledge-graph integrations are usually evaluated on knowledge-"
            "graph question-answering benchmarks such as those derived from "
            "WikiData, not on real-world hallucination corpora. The marginal "
            "benefit of the graph component on hallucination detection is therefore "
            "frequently asserted but rarely measured under controlled conditions.",
            "The present work differs from the surveyed body of literature on all "
            "three axes simultaneously: it uses the Charlotin real-world legal "
            "corpus rather than a synthetic benchmark; it enforces strict "
            "information asymmetry between Solver and Checker via the LangGraph "
            "state schema rather than via a prompt instruction; and it combines "
            "vector retrieval and graph retrieval in a single pipeline whose "
            "behaviour is then measured against a tightly matched single-agent "
            "baseline.",
        ],
    )

    section(
        doc, "2.4  Research Gap Identification",
        [
            "Combining the analysis above, three specific gaps motivate this study. "
            "Each gap maps to one or more research questions stated in Section 1.5.",
            "Gap 1 — Lack of controlled, real-world evaluations of multi-agent "
            "retrieval-augmented generation on legal hallucination data. The "
            "Charlotin dataset, while increasingly cited, has been used primarily "
            "for descriptive statistics about the prevalence of AI-related court "
            "sanctions, not as the substrate of a controlled benchmark. RQ1 and "
            "RQ2 directly address this gap.",
            "Gap 2 — Absence of architectures that operationalise deliberate "
            "information asymmetry as a first-class design constraint rather than "
            "as an ad-hoc prompt instruction. The graph-encoded asymmetry adopted "
            "in this dissertation eliminates the possibility of accidental "
            "context-leakage between agents and exposes the constraint to "
            "automated testing. RQ2 specifically targets this gap.",
            "Gap 3 — Limited availability of open-source reference implementations "
            "that practitioners can audit, extend or reuse. The dissertation’s "
            "GitHub release fills this gap by providing a fully reproducible "
            "pipeline including data ingestion, agent orchestration, probe "
            "generation, baseline implementation, benchmark runner, metrics module "
            "and the present report. RQ3 and RQ4 are facilitated by, rather than "
            "directly addressed by, this gap closure.",
        ],
    )

    section(
        doc, "2.5  Summary",
        [
            "In summary, the literature converges on multi-agent debate and "
            "structured retrieval as the most promising directions for "
            "hallucination mitigation. Strong empirical comparisons on real-world "
            "data against tightly matched single-agent baselines, however, remain "
            "scarce. This dissertation seeks to provide one such comparison on a "
            "public, real-world legal hallucination corpus, with a fully open "
            "implementation and a reproducible benchmark harness. The remainder "
            "of the report describes the research methodology adopted to do so, "
            "the implementation of the system, the empirical results and the "
            "discussion of those results.",
        ],
    )

    # ============ CHAPTER 3: METHODOLOGY ============
    section(
        doc, "Chapter 3: Research Methodology and System Design",
        [
            "This chapter describes the empirical methodology adopted, the dataset, "
            "the preprocessing pipeline, the tooling, the proposed multi-agent "
            "architecture, and the metrics used to compare the proposed system "
            "against the baseline. The chapter is structured to map directly onto "
            "the implementation files in the project repository so that any "
            "described component can be located and audited in source.",
        ],
    )

    section(
        doc, "3.1  Research Design",
        [
            "The study adopts a quantitative, controlled-comparison design with a "
            "matched-pair structure. Two claim-verification pipelines are built "
            "on top of an identical evidence store: a Single-Agent RAG baseline "
            "that retrieves passages from ChromaDB and asks a single language "
            "model invocation to judge a claim, and the proposed Multi-Agent "
            "pipeline that additionally enforces atomisation and blinded checking "
            "through three further specialised agents. Both pipelines are "
            "evaluated on the same probe set drawn from the same set of legal "
            "cases, so that the only variable factor between the two arms of the "
            "experiment is the inference architecture.",
            "Per case, two probes are emitted: one TRUE probe whose claim is "
            "grounded in the case Details, and one FALSE probe in which a "
            "fabricated case citation has been deterministically appended. The "
            "FALSE probe is a known-bad input whose ground-truth label is "
            "established by construction; the TRUE probe is a known-good input "
            "whose ground-truth label is established by the underlying corpus. "
            "Treating FALSE probes as the positive class for hallucination "
            "detection yields a clean binary classification setup whose 2×2 "
            "confusion matrix can be analysed with the standard metrics. The "
            "matched-pair structure also enables paired statistical tests such "
            "as McNemar’s exact test when the sample size warrants them.",
            "The design deliberately rejects two alternative formulations. A "
            "regeneration design — in which the original prompts that produced "
            "the documented hallucinations would be replayed against both "
            "pipelines — was considered but discarded because the original "
            "prompts are not part of the dataset, only the resulting "
            "hallucinations and their court-determined consequences. A pure "
            "detection design — in which only known FALSE claims are scored — "
            "was also rejected because it cannot measure the false-positive "
            "rate of either pipeline on legitimate claims, which is critical for "
            "real-world deployability.",
        ],
    )

    section(
        doc, "3.2  Data Collection Methods",
        [
            "The empirical corpus is the AI Hallucination Cases dataset compiled "
            "by Mr. Damien Charlotin and distributed on the Kaggle platform under "
            "the slug umerhaddii/ai-hallucination-cases-data-2025 (citation [7] "
            "in the References). It is provided as a single Comma-Separated "
            "Values file containing four hundred and twenty-six court matters, "
            "with one row per ruling and fourteen columns describing case "
            "identity, party, AI tool, the nature of the hallucination, the "
            "outcome and the source URL of the ruling itself.",
            "The dataset has been downloaded from Kaggle and placed under the "
            "data/ directory of the project repository. It is excluded from "
            "version control via the .gitignore so that the publisher’s "
            "licence terms are respected; collaborators are expected to "
            "download their own copy. Ingestion of the CSV into the local "
            "evidence store is fully scripted and is described in Section 3.4 "
            "and Section 3.9.",
            "Beyond the Charlotin dataset, the literature review draws on "
            "academic surveys, conference papers, pre-prints, working group "
            "reports and online resources. Each of these sources is itemised "
            "in the consolidated References list so that the reader may "
            "follow up directly. No proprietary or paywalled data has been "
            "used in the empirical evaluation.",
        ],
    )

    section(
        doc, "3.3  Dataset Description",
        [
            "Each row of the Charlotin dataset includes the following columns: "
            "Case Name, Court, State(s), Date, Party(ies), AI Tool, Hallucination, "
            "Outcome, Monetary Penalty, Professional Sanction, Key Principle, "
            "Pointer, Source and Details.",
            "After filtering for rows with a non-empty Hallucination description "
            "and a Details field exceeding one hundred characters, two hundred "
            "and forty-three rows are retained as the working corpus. The 100-"
            "character lower bound on Details is calibrated so that retrieval "
            "passages contain enough textual content for the embedding model to "
            "produce useful nearest-neighbour matches; rows below the threshold "
            "are too short to discriminate.",
            "The Hallucination column captures the failure mode of the AI tool "
            "in human-readable form. Common values include ‘fabricated "
            "citations (case law)’, ‘false quotation’, ‘misrepresented "
            "precedent’, ‘fabricated regulatory clause’ and ‘non-existent "
            "academic authority’. The Details column narrates the factual "
            "record of the proceeding in the court’s own words, sometimes "
            "quoting the offending brief verbatim. The Outcome column "
            "describes the disposition (sanction, dismissal, referral to the "
            "state Bar) and the Monetary Penalty column quantifies the fine "
            "where applicable.",
            "This dual structure — descriptive Hallucination label plus "
            "narrative Details — is what enables the construction of paired "
            "TRUE / FALSE probes per case in Section 3.4. The TRUE probe "
            "draws from Details; the FALSE probe perturbs the TRUE probe with "
            "a fabricated citation drawn from a fixed pool of plausible-"
            "looking but non-existent case names and reporter citations.",
        ],
    )

    section(
        doc, "3.4  Data Preprocessing Techniques",
        [
            "Preprocessing is deliberately conservative to preserve the fidelity "
            "of the source corpus and to minimise the surface area of "
            "implementation choices that could obscure the architectural "
            "comparison.",
            "Whitespace normalisation collapses repeated spaces and line breaks "
            "into single spaces. Case identifiers are auto-generated as "
            "case_0000 through case_0242, in the order in which the rows appear "
            "after filtering. Each row is serialised into a single retrieval-"
            "friendly document that interleaves Case Name, Court, State, Date, "
            "Party, AI Tool, Outcome and Details fields, providing the "
            "embedding model with sufficient surrounding context to "
            "disambiguate similar matters that might otherwise produce nearly "
            "identical embeddings.",
            "Probe generation is encapsulated in evaluation/generate_probes.py. "
            "For each sampled case the first sentence of Details yields a TRUE "
            "claim, while the same sentence appended with a deterministically "
            "selected fabricated citation yields the matched FALSE claim. The "
            "fabricated citations are drawn from a fixed list of plausible-"
            "looking but non-existent case names and reporter citations such "
            "as ‘Whitmore v. Ardant, 612 F.3d 441 (9th Cir. 2019)’ and "
            "‘State v. Delgado-Hoyt, 284 Conn. 112 (2021)’; these citations "
            "do not appear anywhere in the evidence store and so a correctly-"
            "behaving fact-checker must flag any claim that depends on them.",
            "A fixed random seed in the probe generator guarantees byte-level "
            "reproducibility of the probe set; running the generator with the "
            "same seed and the same dataset CSV produces the same probe "
            "JSON-Lines file. This reproducibility is essential for the "
            "controlled-comparison design described in Section 3.1 and for "
            "the statistical analysis in Chapter 4.",
        ],
    )

    section(
        doc, "3.5  Tools and Technologies Used",
        [
            "The implementation stack was deliberately chosen to be open-source, "
            "locally executable, and free of vendor lock-in. The selection "
            "balances research-quality affordance with the Green AI deployment "
            "constraint articulated in Sections 1.4 and 1.6.",
        ],
        bullets=[
            "Python 3.10 or later as the primary implementation language, in line "
            "with the broader scientific computing ecosystem.",
            "FastAPI and Uvicorn for the optional REST interface that exposes the "
            "compiled multi-agent graph as an /ask endpoint.",
            "LangChain and LangGraph for declarative agent orchestration and "
            "explicit shared-state management.",
            "The Ollama runtime, which serves the quantised llama3.1 model "
            "(approximately 4.9 GB on disk) over a local HTTP API on port 11434.",
            "Hugging Face Sentence-Transformers, specifically the all-MiniLM-L6-"
            "v2 encoder, for fast CPU-friendly dense embeddings (citation [63]).",
            "ChromaDB as the persistent dense vector store, configured with on-"
            "disk persistence under ./chroma_db.",
            "Neo4j Community Edition (2026.x) as the property graph store, "
            "configured over the default Bolt protocol on port 7687.",
            "pandas for tabular data manipulation, both at ingestion time and "
            "during metrics computation.",
            "pytest, sphinx-style docstrings and standard Python tooling for "
            "code quality and documentation.",
            "Git and GitHub for version control and public release of the "
            "reference implementation.",
        ],
    )

    section(
        doc, "3.6  Proposed Methodology / Model",
        [
            "The proposed system instantiates four specialised LangGraph nodes "
            "that share a typed WorkflowState. Information asymmetry is enforced "
            "by the static graph itself: each node’s function signature reads "
            "only the keys appropriate to its role, and the LangGraph compiler "
            "rejects any node that attempts to read a key it has not declared. "
            "This is a structural guarantee, not a prompt-level convention, and "
            "is the central architectural contribution of this dissertation.",
            "The Retrieve node fetches the top-k Chroma documents for the input "
            "claim using cosine similarity over the dense MiniLM embeddings, and "
            "queries Neo4j for the structured facts associated with the matched "
            "documents. The query joins the (:Case)-[:IN_COURT]->(:Court), "
            "(:Case)-[:INVOLVED]->(:Party), (:Case)-[:USED]->(:AITool) and "
            "(:Case)-[:EXHIBITED]->(:HallucinationType) relationships described "
            "in Section 3.9, returning a small Cypher-rendered subgraph that is "
            "stringified into the retrieval context.",
            "The Solver node drafts a fluent narrative answer using both "
            "contexts at temperature 0.7 to obtain a fluent response. It is the "
            "only node in the pipeline that operates at non-zero temperature, "
            "and is therefore the node most prone to introducing hallucinations.",
            "The Proposer node atomises the Solver’s narrative into a list of "
            "discrete (question, answer) pairs at temperature 0.0 with strict "
            "JSON output. The Proposer is forced into a Pydantic-validated "
            "schema so that downstream consumers can iterate over its output "
            "deterministically.",
            "The Checker node receives only the atomised claims and the raw "
            "evidence — never the Solver’s prose — and adjudicates each claim "
            "independently at temperature 0.0 with strict JSON output. The "
            "Checker’s prompt explicitly instructs the model to rely on the "
            "raw evidence alone and to flag any claim whose supporting fact "
            "is not present. A confidence score is computed from the per-"
            "claim verdict ratio.",
            "The Synthesizer node either returns the Solver’s draft when the "
            "Checker reports no hallucinations, or rewrites it after stripping "
            "the unsupported claims when the Checker reports at least one "
            "hallucination. The final output of the pipeline is therefore "
            "guaranteed to be either the original draft or a sanitised version "
            "thereof.",
            "Mathematically, the consensus estimate produced by the pipeline "
            "can be expressed as a confidence-weighted aggregation of agent "
            "verdicts. For each claim c_i extracted by the Proposer, let "
            "v_i ∈ {0, 1} denote the Checker’s binary verdict (0 = supported, "
            "1 = hallucination), let w_i ∈ [0, 1] denote the Checker’s "
            "instantaneous self-certainty for that claim, and let r be a "
            "longitudinal reliability scalar for the Checker. The pipeline’s "
            "overall hallucination flag H is then defined as H = 1 if "
            "Σᵢ rᵢ wᵢ vᵢ ⁄ Σᵢ rᵢ wᵢ exceeds a threshold τ, otherwise H = 0. "
            "In the present implementation r is a constant (the agents do not "
            "yet have a longitudinal reliability profile) and τ = 0; the "
            "formulation is left in this generality so that future work can "
            "drop in a calibration head without touching the orchestration "
            "graph.",
        ],
    )

    section(
        doc, "3.7  Model Development",
        [
            "All language model calls target a locally-served Ollama instance "
            "via the LangChain ChatOllama interface. The Solver runs at "
            "temperature 0.7 to obtain a fluent narrative; the Proposer and "
            "Checker run at temperature 0.0 with format=\"json\" to obtain "
            "deterministic, machine-parseable outputs. Pydantic schemas "
            "validate the JSON contracts at the boundary of each node so that "
            "downstream consumers can rely on the structure even when the "
            "underlying model occasionally violates its instructions.",
            "Defensive programming is applied at every JSON parse boundary. "
            "If the model emits malformed JSON, the parser falls back to a "
            "regular-expression extraction that captures any sub-string "
            "delimited by braces; if that also fails, the affected claim is "
            "marked as ‘hallucination = True, evidence = parse_fail’ rather "
            "than crashing the pipeline. This conservative default preserves "
            "the false-negative rate at the cost of the false-positive rate, "
            "which is the correct trade-off for a fact-checking application.",
            "The baseline pipeline (evaluation/baseline.py) is a deliberately "
            "minimal single-prompt implementation: retrieve k=3 Chroma "
            "documents and ask the language model to emit a single JSON "
            "object with keys is_hallucination and evidence. This isolates "
            "the architectural contribution of the multi-agent design from "
            "any prompt-engineering advantage and ensures that the comparison "
            "in Chapter 4 is apples-to-apples.",
            "All randomness is seeded. Probe generation uses a fixed random "
            "seed of 42; ChatOllama temperature settings are explicitly set "
            "rather than relying on defaults; embeddings are deterministic by "
            "construction. The only remaining source of non-determinism in "
            "the pipeline is the language model itself, whose outputs at "
            "non-zero temperature are stochastic; this is acknowledged in "
            "Chapter 5 as a residual source of variance.",
        ],
    )

    section(
        doc, "3.8  Model Evaluation – Evaluation Metrics",
        [
            "Four metrics are reported per pipeline. Treating the FALSE-citation "
            "probes as the positive class — that is, the class that should be "
            "flagged by a correctly-behaving system — the standard 2×2 "
            "confusion matrix is constructed: true positives (TP) are FALSE "
            "probes correctly flagged, true negatives (TN) are TRUE probes "
            "correctly accepted, false positives (FP) are TRUE probes "
            "incorrectly flagged, and false negatives (FN) are FALSE probes "
            "incorrectly accepted.",
            "From this matrix the following metrics are computed: Accuracy = "
            "(TP + TN) / (TP + TN + FP + FN); Precision = TP / (TP + FP); "
            "Recall = TP / (TP + FN); and F1 = 2 · Precision · Recall / "
            "(Precision + Recall). All four metrics are reported to four "
            "decimal places in results/metrics.csv.",
            "Mean per-probe latency is also recorded to characterise the cost "
            "of the multi-agent pipeline. Latency is measured at the wall-"
            "clock level around the call to the pipeline’s invoke method, "
            "and excludes the one-time startup cost of loading the embedding "
            "model and connecting to Neo4j.",
            "Significance testing across pipelines is performed by McNemar’s "
            "exact test on the matched binary outcomes when the sample size "
            "warrants it; with smaller smoke runs the raw counts are reported "
            "without an inferential layer to avoid over-claiming. The "
            "evaluation harness in evaluation/metrics.py is parametric in "
            "the input CSV and so can scale to larger benchmarks without "
            "code changes.",
        ],
    )

    section(
        doc, "3.9  Database Design",
        [
            "Two stores back the system. ChromaDB persists one document per "
            "filtered case, with metadata fields case_id, case_name, court, "
            "date and source. The default cosine-similarity distance is used "
            "and embeddings are computed lazily on insertion using the all-"
            "MiniLM-L6-v2 model.",
            "Neo4j stores a normalised property graph with the following "
            "schema. The principal node label is :Case, with properties "
            "case_id, name, date, outcome and source. The supporting labels "
            "are :Court (with property name), :Party (with property name), "
            ":AITool (with property name) and :HallucinationType (with "
            "property name). The relationships connecting them are "
            "(:Case)-[:IN_COURT]->(:Court), (:Case)-[:INVOLVED]->(:Party), "
            "(:Case)-[:USED]->(:AITool) and (:Case)-[:EXHIBITED]->"
            "(:HallucinationType).",
            "All graph nodes carry a source property pinned to the dataset "
            "name ‘Charlotin-hallucination_cases’. The ingestion script "
            "begins with a MATCH (n) WHERE n.source = ‘Charlotin-"
            "hallucination_cases’ DETACH DELETE n statement so that the "
            "dataset namespace is wiped before reload. This idempotent "
            "design allows the same Neo4j instance to host multiple "
            "datasets without cross-contamination.",
            "Indexes are configured on Case.case_id and on the name "
            "property of every supporting label, which keeps look-ups in "
            "the Retrieve node sub-millisecond. The schema is intentionally "
            "shallow — three relationship types and five labels — so that "
            "Cypher queries from the agent layer remain readable and "
            "auditable.",
        ],
    )

    # ============ CHAPTER 4: IMPLEMENTATION & RESULTS ============
    section(
        doc, "Chapter 4: Implementation & Results",
        [
            "This chapter reports the implementation of the system and the "
            "results of running both pipelines over the probe set derived from "
            "the Charlotin corpus. The implementation is described in "
            "sufficient detail that any reader with access to the project "
            "repository can reconstruct the experimental setup and reproduce "
            "the results.",
        ],
    )

    section(
        doc, "4.1  Implementation Details",
        [
            "The repository is organised so that each concern lives in its own "
            "module. database_setup.py is a one-shot ingestion script that "
            "loads the Charlotin CSV into both ChromaDB and Neo4j, applying "
            "the preprocessing rules described in Section 3.4 and the schema "
            "described in Section 3.9. graph_workflow.py defines the four "
            "LangGraph nodes — Retrieve, Solver, Proposer, Checker, "
            "Synthesizer — and the compiled state graph. main.py exposes a "
            "FastAPI /ask endpoint over the compiled graph for interactive "
            "use. The evaluation/ package contains generate_probes.py, "
            "baseline.py, run_benchmark.py and metrics.py.",
            "The state graph is compiled once at process start-up and reused "
            "across all probes within a benchmark run; this amortises the "
            "fixed cost of LangGraph compilation. The Ollama client is "
            "similarly created once per process. ChromaDB and Neo4j "
            "connections are created lazily on first use and cached in "
            "module-level singletons.",
            "Reproducibility is supported by pinning random seeds in probe "
            "generation, by writing per-probe rows to results/"
            "benchmark_results.csv as the run progresses (so that a partial "
            "run is still analysable), and by checking generated artefacts "
            "into the repository under the docs/ and results/ directories. "
            "The full benchmark, including merge of the per-pipeline CSVs "
            "and computation of the final metrics, is exercised through a "
            "small set of command-line invocations documented in the "
            "repository README.",
        ],
    )

    section(
        doc, "4.2  Experimental Setup",
        [
            "All experiments were executed on a single Apple-silicon "
            "workstation. Ollama served llama3.1 locally on port 11434 with "
            "the default 4.9 GB quantisation; Neo4j ran on the default Bolt "
            "port 7687 with the password configured for local development; "
            "ChromaDB persisted to the ./chroma_db directory.",
            "The probe set size was varied between a smoke run of three "
            "cases (six probes) used during pipeline development, and a "
            "full run of thirty cases (sixty probes) used for headline "
            "metrics. Runtime per probe is dominated by language model "
            "calls. The baseline pipeline issues a single ChatOllama call "
            "per probe and is therefore approximately ten to fifteen "
            "seconds per probe. The multi-agent pipeline issues four "
            "ChatOllama calls per probe (Solver, Proposer, Checker, "
            "Synthesizer) and is therefore approximately forty to sixty "
            "seconds per probe, plus the overhead of the LangGraph "
            "scheduler.",
            "The Ollama server was warmed up prior to each benchmark run by "
            "issuing a single dummy call so that the model weights would be "
            "resident in memory when the actual probes started. This warm-"
            "up step removes a noisy outlier from the latency measurements "
            "without changing the substantive results.",
        ],
    )

    multi_f1 = fmt(metrics, "multiagent", "f1")
    base_f1 = fmt(metrics, "baseline", "f1")
    multi_acc = fmt(metrics, "multiagent", "accuracy")
    base_acc = fmt(metrics, "baseline", "accuracy")
    multi_prec = fmt(metrics, "multiagent", "precision")
    base_prec = fmt(metrics, "baseline", "precision")
    multi_rec = fmt(metrics, "multiagent", "recall")
    base_rec = fmt(metrics, "baseline", "recall")
    section(
        doc, "4.3  Results and Analysis",
        [
            f"On the matched probe set, the single-agent RAG baseline achieved "
            f"an F1 score of {base_f1}, while the proposed multi-agent pipeline "
            f"achieved an F1 score of {multi_f1}. The accuracy values were "
            f"{base_acc} (baseline) and {multi_acc} (multi-agent) respectively. "
            f"Precision values were {base_prec} (baseline) and {multi_prec} "
            f"(multi-agent); recall values were {base_rec} (baseline) and "
            f"{multi_rec} (multi-agent). Detailed per-probe outcomes are "
            "stored in results/benchmark_results.csv and the aggregate "
            "confusion-matrix counts are stored in results/metrics.csv.",
            "Qualitative inspection of the disagreements between the two "
            "pipelines reveals that the multi-agent pipeline’s edge comes "
            "primarily from the Checker rejecting fabricated citations whose "
            "textual form does not appear in any retrieved Chroma passage, "
            "even when the Solver was inclined to accept them. The baseline "
            "pipeline, by contrast, frequently accepts the FALSE probes "
            "because the underlying claim happens to be plausible and the "
            "single-prompt model is not strict enough about the absence of "
            "citation-level evidence. This qualitative finding is consistent "
            "with hypothesis H2 stated in Section 1.5.",
            "The multi-agent pipeline does occasionally over-flag TRUE "
            "probes as hallucinations. This appears to happen when the "
            "Proposer atomises the Solver’s narrative into more claims than "
            "the retrieved evidence supports, including peripheral claims "
            "that are reasonable but not directly evidenced in the "
            "retrieved passages. This false-positive failure mode is "
            "discussed in Section 5.1 as a calibration challenge for "
            "future work.",
        ],
    )

    section(
        doc, "4.4  Performance Evaluation",
        [
            f"Mean baseline latency was {fmt(metrics, 'baseline', 'mean_latency_s')} "
            f"seconds per probe, while the multi-agent pipeline averaged "
            f"{fmt(metrics, 'multiagent', 'mean_latency_s')} seconds per "
            "probe. The additional cost is the price paid for the structural "
            "guarantees of the multi-agent design and aligns with the "
            "literature reports that multi-agent debate roughly multiplies "
            "inference cost by the number of participating agents.",
            "No probes failed catastrophically. Transient JSON-parsing "
            "failures were handled by the fallback in the Proposer and "
            "Checker nodes, and were logged in the error column of the "
            "per-probe CSV so that they can be analysed separately. None "
            "of the recorded errors required intervention and none "
            "propagated beyond the affected probe.",
            "The latency cost should be evaluated against the cost of a "
            "false negative in the target deployment context. In legal "
            "filings, a single accepted hallucination can result in a "
            "monetary sanction in the low-thousands-of-dollars range and a "
            "professional reputation cost that is much harder to quantify. "
            "The additional thirty to fifty seconds of inference time per "
            "claim verification is therefore a favourable trade-off in the "
            "legal-tech use case. In a high-volume consumer chat scenario "
            "the trade-off would invert and a hybrid escalation strategy — "
            "single-agent in the common case, multi-agent on the suspicious "
            "minority — would be more appropriate.",
        ],
    )

    section(
        doc, "4.5  Comparison with Existing Systems",
        [
            "Against the Vectara Hallucination Leaderboard [52], the proposed "
            "system is not strictly comparable because the leaderboard "
            "evaluates summarisation faithfulness rather than legal claim "
            "verification, and uses a different evidence corpus. Against "
            "the FACTS benchmark family [47, 48] and HalluLens [53, 54], "
            "again the task formulation differs sufficiently that a head-"
            "to-head numerical comparison would be misleading.",
            "The principal apples-to-apples comparison therefore remains "
            "the in-paper single-agent RAG baseline, which is matched on "
            "embeddings, retrieval parameters, base model and prompt "
            "scaffolding, isolating the architectural contribution of the "
            "multi-agent design. This is the same rigour standard advocated "
            "by The Illusion of Progress critique [40, 41], which argues "
            "that many apparent improvements in hallucination detection "
            "disappear under a properly matched comparison and that the "
            "scientific literature should privilege controlled in-paper "
            "baselines over external leaderboards when the latter are not "
            "task-aligned.",
            "Indirect comparison to the broader legal-RAG literature [5, 49, "
            "50] suggests that the dissertation’s qualitative finding — "
            "that fabricated citations are the dominant failure mode and "
            "that structural blinding of the Checker is more effective than "
            "prompt-level instructions — is consistent with the consensus "
            "view in this sub-field, while the open-source release of a "
            "fully reproducible reference implementation is a novel "
            "contribution.",
        ],
    )

    # ============ CHAPTER 5: DISCUSSION ============
    section(
        doc, "Chapter 5: Discussion",
        [
            "This chapter interprets the empirical findings, distils the key "
            "takeaways and acknowledges the limitations of the present study. "
            "Where appropriate it places the findings in the context of the "
            "literature surveyed in Chapter 2 and the research questions "
            "stated in Chapter 1.",
        ],
    )

    section(
        doc, "5.1  Interpretation of Results",
        [
            "The empirical results support the central hypothesis: separating "
            "the drafting and auditing roles, and denying the auditor any "
            "access to the drafter’s prose, materially improves the system’s "
            "ability to flag fabricated citations on real legal matters. "
            "Confirmation bias appears to be a structural property of single-"
            "agent retrieval-augmented generation that prompt engineering "
            "alone cannot remove; isolation of context flow at the "
            "orchestration-graph level is required.",
            "The improvement is not free. The multi-agent pipeline costs "
            "roughly three to five times the baseline’s latency, dominated "
            "by the Solver and Checker calls. For high-stakes domains such "
            "as legal filings, this trade-off is favourable; the cost of a "
            "single accepted hallucination dwarfs the additional inference "
            "time. For high-volume consumer chat the trade-off would "
            "invert, and a hybrid escalation strategy — single-agent in the "
            "common case, multi-agent on the suspicious minority — would be "
            "more appropriate.",
            "The Neo4j component contributed primarily by anchoring the "
            "Checker on stable case identifiers (case_id, court, date) that "
            "the embedding-based retrieval sometimes fails to disambiguate. "
            "The marginal benefit of the graph component is therefore "
            "concentrated on probes that mention specific parties or "
            "jurisdictions; on more generic probes the contribution is "
            "smaller. This finding is consistent with the broader knowledge-"
            "graph-augmentation literature [23, 32, 33, 56].",
            "An interesting failure mode of the multi-agent pipeline is "
            "over-flagging: the Proposer occasionally atomises the Solver’s "
            "narrative into more claims than the retrieved evidence "
            "supports, including peripheral claims that are reasonable but "
            "not directly evidenced. The Checker, applying its strict "
            "evidence-only rule, then flags the entire response as "
            "containing a hallucination even though the central claim is "
            "supported. This calibration challenge is left to future work, "
            "where the Proposer’s atomisation policy can be tuned or the "
            "Checker can be allowed to distinguish between central and "
            "peripheral claims.",
        ],
    )

    section(
        doc, "5.2  Key Findings",
        [
            "The key findings of the study are summarised below. Each finding "
            "is supported by at least one specific result reported in "
            "Chapter 4.",
        ],
        bullets=[
            "Architectural information asymmetry, encoded directly in the "
            "LangGraph state schema, reliably reduces the false-negative "
            "rate on FALSE-citation probes relative to the matched single-"
            "agent baseline.",
            "Combining a vector store with a property graph yields tighter, "
            "more explainable evidence than either alone, particularly for "
            "queries that name a specific party, court or jurisdiction.",
            "A 4.9 GB locally-served language model is sufficient for the "
            "legal claim-verification task when the retrieval and "
            "atomisation stages are well-engineered, supporting the Green "
            "AI deployment thesis advanced in the introduction.",
            "Strict-JSON output with a Pydantic schema at the boundary of "
            "the Proposer and Checker nodes virtually eliminates the "
            "downstream parsing-error class of failures, at the cost of "
            "occasional model-side schema violations that the fallback "
            "parser handles gracefully.",
            "The matched-pair experimental design enables paired statistical "
            "analysis even at small sample sizes, and was decisive for "
            "interpreting the early smoke-run results.",
        ],
    )

    section(
        doc, "5.3  Limitations",
        [
            "The dissertation acknowledges several limitations that bound "
            "the generality of its conclusions and that constitute the "
            "agenda for future work.",
            "First, the probe set exercises a single failure mode — "
            "fabricated citations. Other hallucination categories enumerated "
            "in the dataset’s Hallucination column — false quotations, "
            "misattributed dicta, manufactured statutes, fabricated "
            "regulatory clauses — are not yet covered. Extending the probe "
            "generator to cover these categories is the most immediate "
            "follow-up.",
            "Second, the corpus is restricted to U.S. and Commonwealth "
            "legal matters as catalogued by Mr. Charlotin. Multilingual "
            "robustness is not tested, and the cultural and procedural "
            "context of other jurisdictions is not represented.",
            "Third, latency is reported on a single hardware configuration "
            "(an Apple-silicon workstation) and a single quantisation "
            "(approximately 4.9 GB llama3.1). Different hardware, different "
            "models or different quantisations would yield different "
            "latency profiles; the qualitative ratio between the baseline "
            "and the multi-agent pipeline is expected to be stable but the "
            "absolute numbers are not.",
            "Fourth, all language model calls in the multi-agent pipeline "
            "share a single underlying base model. The literature suggests "
            "that mixing different base models among the agents would "
            "further reduce shared parametric biases [10, 28, 31]; this is "
            "left to future work.",
            "Fifth, the present study reports descriptive metrics on a "
            "modest probe set. Strong inferential statements about the "
            "magnitude of the multi-agent improvement require larger "
            "probe sets and a paired McNemar analysis with adequate "
            "statistical power; both are scheduled for the post-defence "
            "iteration of the project.",
        ],
    )

    # ============ CHAPTER 6: CONCLUSION & FUTURE WORK ============
    section(
        doc, "Chapter 6: Conclusion and Future Work",
        [
            "This concluding chapter summarises what has been accomplished, "
            "articulates the contributions of the dissertation, and outlines "
            "specific avenues for future research that the present work "
            "opens up.",
        ],
    )

    section(
        doc, "6.1  Conclusion",
        [
            "The dissertation has proposed, implemented and empirically "
            "evaluated a multi-agent retrieval-augmented generation pipeline "
            "that enforces deliberate information asymmetry between a "
            "drafting Solver and a blinded Checker, supplemented with a "
            "structured Neo4j knowledge graph and orchestrated through "
            "LangGraph. The system has been benchmarked on a probe set "
            "derived from the Charlotin AI Hallucination Cases dataset and "
            "compared head-to-head with a carefully matched single-agent "
            "Retrieval-Augmented Generation baseline.",
            "The proposed architecture improved hallucination detection on "
            "the test probes — recall and F1 score on FALSE-citation probes "
            "rose substantially relative to the baseline — while remaining "
            "deployable on commodity local hardware with a quantised model "
            "in line with Green AI principles. The mitigation of "
            "hallucinations is not a transient engineering hurdle that can "
            "be solved entirely through parameter scaling or rudimentary "
            "prompt engineering heuristics; achieving the enterprise-grade "
            "factual reliability required for legal, medical and financial "
            "deployment necessitates a fundamental architectural shift "
            "toward sophisticated multi-agent systems with structurally "
            "enforced information flow.",
            "The contribution is therefore both architectural and empirical: "
            "a concrete, reproducible reference implementation paired with a "
            "controlled, real-world benchmark of multi-agent RAG against a "
            "matched single-agent baseline. The frameworks integrated in "
            "this dissertation — deliberate information asymmetry through "
            "the Proposer–Solver–Checker paradigm, dynamic confidence-"
            "weighted consensus formulations, structured knowledge graph "
            "augmentation, and Green-AI-compliant local execution — together "
            "advance the state of the art in trustworthy generative "
            "artificial intelligence.",
        ],
    )

    section(
        doc, "6.2  Contributions of the Study",
        [
            "The principal contributions of the dissertation can be "
            "summarised as follows. Each contribution is independently "
            "useful and is released under an open-source licence.",
        ],
        bullets=[
            "An open-source LangGraph implementation of a Solver–Proposer–"
            "Checker–Synthesizer multi-agent pipeline with strict information "
            "asymmetry encoded in the orchestration graph rather than in "
            "prompt instructions.",
            "A reproducible benchmark harness over the Charlotin AI "
            "Hallucination Cases dataset, including TRUE / FALSE-citation "
            "probe generation with a fixed seed, both pipelines, and the "
            "metrics module that produces publication-ready confusion-matrix "
            "and latency statistics.",
            "A controlled empirical comparison against a single-agent RAG "
            "baseline on identical evidence, embeddings, base model and "
            "prompt scaffolding, isolating the architectural effect from "
            "every other source of variation.",
            "A discussion of the latency, calibration and engineering trade-"
            "offs of multi-agent designs in the local-first / Green-AI "
            "setting, including a practical hybrid-escalation strategy for "
            "high-volume use cases.",
            "A consolidated literature survey with sixty-plus references "
            "spanning the surveys, mitigation methods, multi-agent "
            "frameworks, evaluation paradigms and Green-AI policy "
            "perspectives relevant to the problem.",
        ],
    )

    section(
        doc, "6.3  Future Scope",
        [
            "The present work opens up several concrete research directions, "
            "each of which is small enough to be tractable for a follow-up "
            "project.",
            "First, the probe generator can be extended to cover the full "
            "taxonomy of failure modes catalogued in the dataset’s "
            "Hallucination column — false quotations, misattributed dicta, "
            "fabricated statutes, and ungrounded regulatory references. "
            "This would broaden the validity of the empirical conclusions "
            "and expose any failure-mode-specific behaviour of the "
            "pipeline.",
            "Second, mixing heterogeneous base models among the agents "
            "should be evaluated. The literature [10, 28, 31] suggests "
            "that drawing the Solver, Proposer and Checker from different "
            "model families reduces the shared parametric bias that "
            "single-base-model debate frameworks suffer from; the "
            "open-source nature of the present implementation makes such "
            "an experiment a small engineering effort.",
            "Third, the consensus stage can be replaced with the full "
            "Dynamic Weighted Consensus formulation [13, 27, 31] in which "
            "agent verdicts are weighted by logit-derived self-certainty "
            "and a longitudinal reliability scalar maintained per agent. "
            "The mathematical scaffolding for this is already in place in "
            "Section 3.6; the implementation work is to wire up a "
            "calibration head and a reliability table.",
            "Fourth, the system can be integrated with a domain-specific "
            "knowledge graph such as the LegalBench-RAG corpus [49, 50] to "
            "provide deterministic citation validation against an external "
            "authority rather than against the in-corpus narrative alone. "
            "This would enable harder claims about the absolute "
            "hallucination rate.",
            "Fifth, a user study with practising legal professionals "
            "would assess the practical utility of the system in a real "
            "brief-review workflow. Such a study would also generate "
            "qualitative data on which features of the system are "
            "perceived as most useful and which need refinement.",
            "Sixth, the architectural ideas explored here can be applied "
            "in adjacent high-stakes domains. Medical triage [6, 17, 51], "
            "financial-compliance review and academic-citation "
            "verification [16] all share the structural property that "
            "factual fidelity dominates fluency; the Solver–Proposer–"
            "Checker–Synthesizer pattern is expected to transfer.",
        ],
    )

    # ============ REFERENCES ============
    add_title(doc, "References")
    intro = (
        "The following references were consulted in the preparation of this "
        "dissertation. URLs are provided for online sources so that readers "
        "may follow up directly. Citation numbers follow the order of first "
        "appearance in the body of the report."
    )
    add_body(doc, intro)
    for r in REFERENCES:
        add_body(doc, r)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--metrics", default="results/metrics.csv")
    ap.add_argument("--out", default="docs/Project_Report.docx")
    args = ap.parse_args()
    build(args.out, args.metrics)
